#!/usr/bin/env python3
"""
docx_export.py — Генерация DOCX/TXT отчётов для чтеца

Выделен из web_viewer.py для отдельного использования.
Генерирует отформатированные отчёты об ошибках чтеца.

Форматирование:
    КАПС КУРСИВ — ошибка (что сказал чтец / услышал Яндекс)
    жирный — правильное слово из текста
    (**слово**) — пропущенное слово
    ⇄ — маркер перестановки слов

Версия: 5.1.0
"""

VERSION = '5.1.0'
VERSION_DATE = '2026-01-25'

import os
from pathlib import Path
from typing import List, Dict, Any, Optional


# =============================================================================
# МЕТКИ ТИПОВ ОШИБОК
# =============================================================================

TYPE_LABELS = {
    'substitution': 'Замена',
    'transposition': 'Перестановка',
    'insertion': 'Лишнее слово',
    'deletion': 'Пропуск',
}


# =============================================================================
# ГЕНЕРАЦИЯ DOCX
# =============================================================================

def generate_reader_docx(
    errors: List[Dict[str, Any]],
    output_dir: str,
    chapter_number: str = '1',
    filename: Optional[str] = None
) -> str:
    """
    Генерирует DOCX файл для чтеца.

    Args:
        errors: список ошибок (dict с полями type, time_formatted, wrong, correct, word, context)
        output_dir: папка для сохранения
        chapter_number: номер главы
        filename: имя файла (None = авто)

    Returns:
        Имя созданного файла
    """
    os.makedirs(output_dir, exist_ok=True)

    if filename is None:
        filename = f'{chapter_number}_для_чтеца.docx'
    filepath = os.path.join(output_dir, filename)

    try:
        return _generate_docx(errors, filepath, chapter_number)
    except ImportError:
        return _generate_txt_fallback(errors, output_dir, chapter_number)


def _generate_docx(
    errors: List[Dict[str, Any]],
    filepath: str,
    chapter_number: str
) -> str:
    """Генерация DOCX через python-docx."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # === ЛЕГЕНДА ФОРМАТИРОВАНИЯ ===
    legend_title = doc.add_paragraph()
    run = legend_title.add_run('Легенда форматирования:')
    run.bold = True

    legend_items = [
        '• КАПС КУРСИВ — ошибка (что сказал чтец / услышал Яндекс)',
        '• жирный — правильное слово из текста',
        '• (**слово**) — пропущенное слово',
        '• ⇄ — маркер перестановки слов',
    ]
    for item in legend_items:
        doc.add_paragraph(item)

    # Разделитель
    separator = doc.add_paragraph()
    separator.add_run('_' * 50)

    # === ГЛАВА ===
    chapter_para = doc.add_paragraph()
    run = chapter_para.add_run(f'Глава {chapter_number}')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()  # пустая строка

    # === ОШИБКИ ===
    for error in errors:
        _add_error_to_docx(doc, error)

    doc.save(filepath)
    print(f"  -> DOCX для чтеца создан: {filepath}")
    return os.path.basename(filepath)


def _add_error_to_docx(doc, error: Dict[str, Any]) -> None:
    """Добавляет одну ошибку в DOCX документ."""
    time_str = error.get('time_formatted', '0:00')
    context = error.get('context', '')
    error_type = error.get('type', '')

    # --- Тип ошибки ---
    type_para = doc.add_paragraph()
    run_type = type_para.add_run(TYPE_LABELS.get(error_type, error_type))
    run_type.italic = True

    # --- Строка с таймкодом ---
    p = doc.add_paragraph()
    run_time = p.add_run(f'{time_str}')
    run_time.bold = True

    if error_type == 'substitution':
        wrong = error.get('wrong', '')
        correct = error.get('correct', '')
        p.add_run(' — ')
        run_wrong = p.add_run(wrong.upper())
        run_wrong.italic = True
        p.add_run(' → ')
        run_correct = p.add_run(correct)
        run_correct.bold = True

    elif error_type == 'transposition':
        wrong = error.get('wrong', '')
        correct = error.get('correct', '')
        p.add_run(' — ')
        run_wrong = p.add_run(wrong.upper())
        run_wrong.italic = True
        p.add_run(' ⇄ ')
        run_correct = p.add_run(correct)
        run_correct.bold = True

    elif error_type == 'insertion':
        word = error.get('word', '')
        p.add_run(' — лишнее слово: ')
        run_word = p.add_run(word.upper())
        run_word.italic = True

    elif error_type == 'deletion':
        word = error.get('word', '')
        p.add_run(' — пропущено: ')
        run_word = p.add_run(word)
        run_word.bold = True

    # --- Контекст с выделением ---
    if context:
        ctx_para = doc.add_paragraph()
        ctx_para.add_run('...')
        _format_context_in_docx(ctx_para, error, context)
        ctx_para.add_run('...')

    # Пустая строка между ошибками
    doc.add_paragraph()


def _format_context_in_docx(ctx_para, error: Dict[str, Any], context: str) -> None:
    """Форматирует контекст ошибки с выделением слов."""
    error_type = error.get('type', '')

    if error_type == 'substitution':
        correct = error.get('correct', '')
        wrong = error.get('wrong', '')
        context_lower = context.lower()
        correct_lower = correct.lower()

        if correct_lower in context_lower:
            idx = context_lower.index(correct_lower)
            before = context[:idx]
            word_in_context = context[idx:idx + len(correct)]
            after = context[idx + len(correct):]

            ctx_para.add_run(before)
            run_w = ctx_para.add_run(wrong.upper())
            run_w.italic = True
            ctx_para.add_run(' (')
            run_c = ctx_para.add_run(word_in_context)
            run_c.bold = True
            ctx_para.add_run(')')
            ctx_para.add_run(after)
        else:
            ctx_para.add_run(context)

    elif error_type == 'deletion':
        word = error.get('word', '')
        context_lower = context.lower()
        word_lower = word.lower()

        if word_lower in context_lower:
            idx = context_lower.index(word_lower)
            before = context[:idx]
            word_in_context = context[idx:idx + len(word)]
            after = context[idx + len(word):]

            ctx_para.add_run(before)
            ctx_para.add_run('(**')
            run_word = ctx_para.add_run(word_in_context)
            run_word.bold = True
            ctx_para.add_run('**)')
            ctx_para.add_run(after)
        else:
            ctx_para.add_run(context)
            ctx_para.add_run(' [ПРОПУЩЕНО: **')
            run_w = ctx_para.add_run(word)
            run_w.bold = True
            ctx_para.add_run('**]')

    elif error_type == 'insertion':
        word = error.get('word', '')
        context_lower = context.lower()
        word_lower = word.lower()

        if word_lower in context_lower:
            idx = context_lower.index(word_lower)
            before = context[:idx]
            word_in_context = context[idx:idx + len(word)]
            after = context[idx + len(word):]

            ctx_para.add_run(before)
            run_w = ctx_para.add_run(word_in_context.upper())
            run_w.italic = True
            ctx_para.add_run(after)
        else:
            ctx_para.add_run(context)

    elif error_type == 'transposition':
        correct = error.get('correct', '')
        wrong = error.get('wrong', '')
        context_lower = context.lower()
        correct_lower = correct.lower()

        if correct_lower in context_lower:
            idx = context_lower.index(correct_lower)
            before = context[:idx]
            phrase_in_context = context[idx:idx + len(correct)]
            after = context[idx + len(correct):]

            ctx_para.add_run(before)
            run_marker = ctx_para.add_run('⇄')
            run_marker.bold = True
            run_w = ctx_para.add_run(wrong.upper())
            run_w.italic = True
            ctx_para.add_run(' (')
            run_c = ctx_para.add_run(phrase_in_context)
            run_c.bold = True
            ctx_para.add_run(')')
            ctx_para.add_run(after)
        else:
            ctx_para.add_run(context)

    else:
        ctx_para.add_run(context)


# =============================================================================
# TXT FALLBACK
# =============================================================================

def _generate_txt_fallback(
    errors: List[Dict[str, Any]],
    output_dir: str,
    chapter_number: str
) -> str:
    """Генерация TXT если python-docx недоступен."""
    txt_filename = f'Чтецу_глава_{chapter_number}.txt'
    txt_filepath = os.path.join(output_dir, txt_filename)

    with open(txt_filepath, 'w', encoding='utf-8') as f:
        f.write('Легенда форматирования:\n')
        f.write('• КАПС КУРСИВ — ошибка (что сказал чтец / услышал Яндекс)\n')
        f.write('• жирный — правильное слово из текста\n')
        f.write('• (**слово**) — пропущенное слово\n')
        f.write('• ⇄ — маркер перестановки слов\n')
        f.write('\n' + '_' * 50 + '\n\n')
        f.write(f'Глава {chapter_number}\n\n')

        for error in errors:
            time_str = error.get('time_formatted', '0:00')
            context = error.get('context', '')
            error_type = error.get('type', '')

            f.write(f'{TYPE_LABELS.get(error_type, error_type)}\n')

            if error_type == 'substitution':
                wrong = error.get('wrong', '')
                correct = error.get('correct', '')
                f.write(f'{time_str} — {wrong.upper()} → {correct}\n')
            elif error_type == 'transposition':
                wrong = error.get('wrong', '')
                correct = error.get('correct', '')
                f.write(f'{time_str} — {wrong.upper()} ⇄ {correct}\n')
            elif error_type == 'deletion':
                word = error.get('word', '')
                f.write(f'{time_str} — пропущено: {word}\n')
            elif error_type == 'insertion':
                word = error.get('word', '')
                f.write(f'{time_str} — лишнее: {word.upper()}\n')

            if context:
                f.write(f'...{context}...\n')
            f.write('\n')

    print(f"  -> TXT для чтеца создан: {txt_filepath}")
    return txt_filename
