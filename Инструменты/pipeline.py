#!/usr/bin/env python3
"""
Pipeline v2.1 - Полный пайплайн проверки аудиокниги

Объединяет все этапы:
1. Конвертация аудио в OggOpus
2. Отправка в Яндекс SpeechKit
3. Нормализация оригинального текста
4. Умное сравнение (якоря + серые зоны)
5. Фильтрация через золотой стандарт
6. Генерация отчёта для чтеца

Использование:
    python pipeline.py глава.mp3 оригинал.docx
    python pipeline.py глава.mp3 оригинал.txt --skip-transcribe  # если уже есть JSON
    python pipeline.py глава.mp3 оригинал.docx --web  # открыть в браузере
    python pipeline.py глава.mp3 оригинал.docx --force  # перезаписать результаты
    python pipeline.py глава.mp3 оригинал.docx --verbose  # подробный вывод
    python pipeline.py глава.mp3 оригинал.docx --quiet  # только ошибки

Changelog:
    v2.1 (2026-01-24): Интеграция системы логирования
        - setup_logging(), get_logger() из config.py
        - Флаги --verbose/--quiet для управления уровнем логирования
        - Логирование всех этапов пайплайна
        - Автоматическая ротация логов
    v2.0 (2026-01-24): Полная интеграция с config.py
        - YandexCloudConfig для folder_id и API ключей
        - RESULTS_DIR, DICTIONARIES_DIR, READER_ERRORS
        - FileNaming для всех выходных файлов
        - SmartCompareConfig для threshold/phantom
        - check_file_exists() + флаг --force
        - VERSION/VERSION_DATE константы
    v1.0: Базовая версия пайплайна
"""

# Версия модуля
VERSION = '5.0.0'
VERSION_DATE = '2026-01-25'

import argparse
import json
import os
import sys
import subprocess
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple

# Опционально импортируем tqdm для прогресс-бара
try:
    from tqdm import tqdm
    HAS_TQDM = True
except ImportError:
    HAS_TQDM = False
    tqdm = None


# Добавляем путь к модулям
SCRIPT_DIR = Path(__file__).parent
sys.path.insert(0, str(SCRIPT_DIR))

# =============================================================================
# ИМПОРТ ЦЕНТРАЛИЗОВАННОЙ КОНФИГУРАЦИИ
# =============================================================================

try:
    from config import (
        PROJECT_DIR, RESULTS_DIR, DICTIONARIES_DIR,
        READER_ERRORS, NAMES_DICT,
        FileNaming, YandexCloudConfig, SmartCompareConfig,
        check_file_exists, validate_file_correspondence,
        setup_logging, get_logger, cleanup_old_logs, log_exception
    )
    HAS_CONFIG = True
    HAS_LOGGING = True
except ImportError:
    HAS_CONFIG = False
    HAS_LOGGING = False
    # Fallback значения
    PROJECT_DIR = SCRIPT_DIR.parent
    RESULTS_DIR = PROJECT_DIR / 'Результаты проверки'
    DICTIONARIES_DIR = PROJECT_DIR / 'Словари'
    READER_ERRORS = DICTIONARIES_DIR / 'ошибки_чтеца.json'
    NAMES_DICT = DICTIONARIES_DIR / 'Словарь_имён_персонажей.txt'

    def check_file_exists(path, action='skip'):
        """Fallback проверка существования файла."""
        if not path.exists():
            return True
        if action == 'overwrite':
            return True
        print(f"  → Файл уже существует: {path.name}")
        return action != 'skip'

# Fallback для логирования если config.py недоступен
if not HAS_LOGGING:
    import logging

    def setup_logging(level=None, log_file=None, module_name=None, console=True, session_id=None):
        """Fallback настройка логирования."""
        logger = logging.getLogger(module_name or 'pipeline')
        if not logger.handlers:
            handler = logging.StreamHandler()
            handler.setFormatter(logging.Formatter('%(asctime)s [%(levelname)s] %(name)s: %(message)s'))
            logger.addHandler(handler)
        logger.setLevel(getattr(logging, (level or 'INFO').upper()))
        return logger

    def get_logger(name=None):
        """Fallback получение логгера."""
        return logging.getLogger(name or 'pipeline')

    def cleanup_old_logs(logs_dir, keep_count=10):
        """Fallback ротация логов."""
        return 0

    def log_exception(logger, msg, exc=None):
        """Fallback логирование исключений."""
        if exc:
            logger.error(f"{msg}: {exc}", exc_info=True)
        else:
            logger.error(msg, exc_info=True)


# Инициализируем логгер модуля (будет настроен в main)
logger = get_logger('pipeline')


# =============================================================================
# ВАЛИДАЦИЯ JSON ТРАНСКРИПЦИИ
# =============================================================================

class TranscriptValidationError(Exception):
    """Ошибка валидации транскрипции"""
    def __init__(self, message: str, details: Optional[Dict] = None):
        super().__init__(message)
        self.details = details or {}


def validate_transcript_json(data: Dict[str, Any]) -> Tuple[bool, List[str]]:
    """
    Валидирует структуру JSON транскрипции от Яндекс SpeechKit.

    Args:
        data: JSON данные транскрипции

    Returns:
        Кортеж (is_valid, errors) где errors — список ошибок
    """
    errors = []

    # Проверяем наличие ключевых полей
    if not isinstance(data, dict):
        errors.append("JSON должен быть объектом (dict)")
        return False, errors

    # Проверяем chunks
    if 'chunks' not in data:
        errors.append("Отсутствует поле 'chunks'")
    elif not isinstance(data['chunks'], list):
        errors.append("Поле 'chunks' должно быть списком")
    elif len(data['chunks']) == 0:
        errors.append("Список 'chunks' пуст — транскрипция не содержит данных")
    else:
        # Валидируем структуру chunks
        for i, chunk in enumerate(data['chunks'][:5]):  # Проверяем первые 5
            if not isinstance(chunk, dict):
                errors.append(f"chunks[{i}] должен быть объектом")
                continue

            if 'alternatives' not in chunk:
                errors.append(f"chunks[{i}] не содержит 'alternatives'")
            elif not isinstance(chunk['alternatives'], list):
                errors.append(f"chunks[{i}]['alternatives'] должен быть списком")
            elif len(chunk['alternatives']) > 0:
                alt = chunk['alternatives'][0]
                if not isinstance(alt, dict):
                    errors.append(f"chunks[{i}]['alternatives'][0] должен быть объектом")
                elif 'words' not in alt:
                    errors.append(f"chunks[{i}]['alternatives'][0] не содержит 'words'")
                elif 'text' not in alt:
                    errors.append(f"chunks[{i}]['alternatives'][0] не содержит 'text'")

    # Проверяем общую статистику
    if 'chunks' in data and isinstance(data['chunks'], list):
        total_words = 0
        total_text_length = 0

        for chunk in data['chunks']:
            if isinstance(chunk, dict) and 'alternatives' in chunk:
                for alt in chunk.get('alternatives', []):
                    if isinstance(alt, dict):
                        total_words += len(alt.get('words', []))
                        total_text_length += len(alt.get('text', ''))

        if total_words == 0:
            errors.append("Транскрипция не содержит слов (words пустые)")

        if total_text_length == 0:
            errors.append("Транскрипция не содержит текста (text пустой)")

    return len(errors) == 0, errors


def validate_transcript_file(file_path: str) -> Dict[str, Any]:
    """
    Валидирует файл транскрипции и возвращает данные или выбрасывает исключение.

    Args:
        file_path: путь к JSON файлу

    Returns:
        Валидные данные транскрипции

    Raises:
        TranscriptValidationError: если валидация не прошла
    """
    path = Path(file_path)

    # Проверяем существование файла
    if not path.exists():
        raise TranscriptValidationError(
            f"Файл транскрипции не найден: {file_path}"
        )

    # Проверяем расширение
    if path.suffix.lower() != '.json':
        raise TranscriptValidationError(
            f"Файл должен иметь расширение .json: {file_path}"
        )

    # Проверяем размер (минимум 100 байт для валидной транскрипции)
    if path.stat().st_size < 100:
        raise TranscriptValidationError(
            f"Файл слишком маленький ({path.stat().st_size} байт): {file_path}"
        )

    # Загружаем JSON
    try:
        with open(path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        raise TranscriptValidationError(
            f"Ошибка парсинга JSON: {e}",
            details={'line': e.lineno, 'column': e.colno}
        )

    # Валидируем структуру
    is_valid, errors = validate_transcript_json(data)

    if not is_valid:
        raise TranscriptValidationError(
            f"Невалидная структура транскрипции ({len(errors)} ошибок)",
            details={'errors': errors}
        )

    print(f"  ✓ Транскрипция валидна")

    # Выводим статистику
    total_chunks = len(data.get('chunks', []))
    total_words = sum(
        len(alt.get('words', []))
        for chunk in data.get('chunks', [])
        for alt in chunk.get('alternatives', [])
    )
    print(f"    Чанков: {total_chunks}, слов: {total_words}")

    return data


class PipelineProgress:
    """Управляет прогресс-баром пайплайна"""

    STEPS = [
        "Конвертация в OggOpus",
        "Транскрибация (Яндекс SpeechKit)",
        "Проверка соответствия файлов",
        "Нормализация текста",
        "Умное сравнение",
        "Фильтр отсева",
        "Генерация DOCX",
    ]

    def __init__(self, use_tqdm: bool = True, skip_steps: List[str] = None):
        self.use_tqdm = use_tqdm and HAS_TQDM
        self.skip_steps = skip_steps or []
        self.pbar = None
        self.current_step = 0

        # Подсчитываем активные шаги
        active_steps = [s for s in self.STEPS if s not in self.skip_steps]
        self.total_steps = len(active_steps)

        if self.use_tqdm:
            self.pbar = tqdm(
                total=self.total_steps,
                desc="Пайплайн",
                unit="шаг",
                bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}]"
            )

    def start_step(self, name: str):
        """Начинает шаг"""
        if self.use_tqdm and self.pbar:
            self.pbar.set_postfix_str(name[:30])
        else:
            print(f"\n{'='*60}")
            print(f"  ШАГ: {name}")
            print(f"{'='*60}")

    def complete_step(self, name: str, success: bool = True):
        """Завершает шаг"""
        if self.use_tqdm and self.pbar:
            self.pbar.update(1)
            if success:
                self.pbar.set_postfix_str(f"✓ {name[:25]}")
            else:
                self.pbar.set_postfix_str(f"✗ {name[:25]}")
        else:
            if success:
                print(f"  ✓ {name} — успешно")
            else:
                print(f"  ✗ {name} — ошибка")

        self.current_step += 1

    def close(self):
        """Закрывает прогресс-бар"""
        if self.pbar:
            self.pbar.close()


# Глобальный прогресс (None если не используется)
_progress: Optional[PipelineProgress] = None


def run_step(name, func, *args, progress: PipelineProgress = None, **kwargs):
    """Выполняет шаг пайплайна с логированием и прогрессом"""
    global _progress
    prog = progress or _progress

    logger.info(f"Начинаю: {name}")

    if prog:
        prog.start_step(name)
    else:
        print(f"\n{'='*60}")
        print(f"  ШАГ: {name}")
        print(f"{'='*60}")

    try:
        result = func(*args, **kwargs)
        logger.info(f"Завершено: {name}")
        if prog:
            prog.complete_step(name, success=True)
        else:
            print(f"  ✓ {name} — успешно")
        return result
    except Exception as e:
        logger.error(f"Ошибка в шаге '{name}': {e}")
        log_exception(logger, f"Детали ошибки в шаге '{name}'", e)
        if prog:
            prog.complete_step(name, success=False)
        else:
            print(f"  ✗ {name} — ошибка: {e}")
        raise


def step_convert_audio(audio_path, output_dir):
    """Шаг 1: Конвертация аудио в OggOpus"""
    from audio_converter import convert_audio, get_audio_info

    audio_path = Path(audio_path)

    # Проверяем, нужна ли конвертация
    if audio_path.suffix.lower() in ('.ogg', '.opus'):
        info = get_audio_info(str(audio_path))
        if info and info.get('channels', 2) == 1:
            print(f"  Файл уже в OggOpus моно, пропускаем конвертацию")
            return str(audio_path)

    # Конвертируем
    output_path = Path(output_dir) / (audio_path.stem + '_yandex.ogg')
    convert_audio(str(audio_path), str(output_path), format='ogg', verbose=True)

    return str(output_path)


def step_transcribe(audio_path, output_dir, api_key=None, folder_id=None):
    """Шаг 2: Транскрибация через Яндекс SpeechKit"""
    from transcribe import transcribe, get_api_key

    if not api_key:
        api_key = get_api_key()

    # Загружаем folder_id из централизованной конфигурации
    if not folder_id:
        if HAS_CONFIG:
            folder_id = YandexCloudConfig.get_folder_id()
        else:
            # Fallback: загрузка из api_keys.json
            keys_file = PROJECT_DIR / 'api_keys.json'
            if keys_file.exists():
                with open(keys_file, 'r', encoding='utf-8') as f:
                    keys = json.load(f)
                    folder_id = keys.get('folder_id')

    audio_path = Path(audio_path)

    # Используем FileNaming для имени файла
    if HAS_CONFIG:
        chapter_id = FileNaming.get_chapter_id(audio_path)
        output_path = Path(output_dir) / FileNaming.build_filename(chapter_id, 'transcript')
    else:
        output_path = Path(output_dir) / (audio_path.stem + '_transcript.json')

    result = transcribe(
        str(audio_path),
        api_key=api_key,
        output_path=str(output_path),
        folder_id=folder_id
    )

    return str(output_path)


def step_normalize_text(text_path, output_dir):
    """Шаг 3: Нормализация оригинального текста"""
    from text_normalizer import normalize_file

    text_path = Path(text_path)
    output_path = Path(output_dir) / (text_path.stem + '_normalized.txt')

    normalize_file(
        str(text_path),
        output_path=str(output_path),
        expand_numbers=True,
        expand_abbrev=True,
        keep_hyphens=False
    )

    return str(output_path)


def step_validate_correspondence(transcript_path, text_path):
    """Шаг 3.5: Проверка соответствия транскрипции и текста"""
    # Используем validate_file_correspondence из config.py (уже импортирован)
    if not HAS_CONFIG:
        print("  ⚠ config.py недоступен, пропускаем проверку соответствия")
        return True

    result = validate_file_correspondence(Path(transcript_path), Path(text_path))

    if result['error']:
        print(f"  ⚠ Ошибка проверки соответствия: {result['error']}")
    elif not result['valid']:
        print(f"  ⚠ ВНИМАНИЕ: транскрипция может не соответствовать тексту!")
        print(f"    Совпадающих слов в начале: {result['matching_start_words']}/20")
        print(f"    Проверьте, что файлы от одной главы.")
    else:
        print(f"  ✓ Файлы соответствуют (совпадений: {result['matching_start_words']}/20)")

    return result['valid']


def step_smart_compare(transcript_path, original_path, audio_path, output_dir, threshold=None, phantom=None):
    """Шаг 4: Умное сравнение"""
    from smart_compare import smart_compare

    # Значения по умолчанию из конфигурации
    if threshold is None:
        threshold = SmartCompareConfig.THRESHOLD if HAS_CONFIG else 0.7
    if phantom is None:
        phantom = SmartCompareConfig.PHANTOM_SECONDS if HAS_CONFIG else -1

    transcript_path = Path(transcript_path)

    # Используем FileNaming для имени файла
    if HAS_CONFIG:
        chapter_id = FileNaming.get_chapter_id(transcript_path)
        output_path = Path(output_dir) / FileNaming.build_filename(chapter_id, 'compared')
    else:
        output_path = Path(output_dir) / (transcript_path.stem.replace('_transcript', '') + '_compared.json')

    report = smart_compare(
        str(transcript_path),
        str(original_path),
        audio_path=str(audio_path),
        threshold=threshold,
        output_path=str(output_path),
        phantom_seconds=phantom
    )

    return str(output_path)


def step_golden_filter(report_path, output_dir, config=None):
    """Шаг 5: Фильтрация через золотой стандарт"""
    import importlib
    import golden_filter

    # Перезагружаем модуль для актуальных словарей
    importlib.reload(golden_filter)
    from golden_filter import filter_report, load_reader_errors

    # Загружаем словари из централизованных путей
    if READER_ERRORS.exists():
        load_reader_errors(str(READER_ERRORS))

    report_path = Path(report_path)

    # Используем FileNaming для имени файла
    if HAS_CONFIG:
        chapter_id = FileNaming.get_chapter_id(report_path)
        output_path = Path(output_dir) / FileNaming.build_filename(chapter_id, 'filtered')
    else:
        output_path = Path(output_dir) / (report_path.stem.replace('_compared', '') + '_filtered.json')

    config = config or {
        'levenshtein_threshold': 2,
        'use_lemmatization': True,
        'use_homophones': True,
    }

    # v14.9: Добавляем chapter в config для cluster_analyzer
    if HAS_CONFIG and chapter_id:
        # chapter_id формата "01", "02", etc. → int
        try:
            config['chapter'] = int(chapter_id)
        except ValueError:
            pass

    filter_report(
        str(report_path),
        output_path=str(output_path),
        **config
    )

    return str(output_path)


def step_generate_docx(report_path, audio_path, output_dir):
    """Шаг 6: Генерация DOCX для чтеца"""
    # Импортируем функцию генерации из web_viewer если есть
    # Или используем простую генерацию

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  ⚠ python-docx не установлен, пропускаем генерацию DOCX")
        return None

    with open(report_path, 'r', encoding='utf-8') as f:
        report = json.load(f)

    errors = report.get('errors', [])
    if not errors:
        print("  Нет ошибок для отчёта")
        return None

    # Создаём документ
    doc = Document()

    # Заголовок
    title = doc.add_heading('Ошибки чтеца', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Информация
    audio_name = Path(audio_path).name if audio_path else 'неизвестно'
    doc.add_paragraph(f'Аудио: {audio_name}')
    doc.add_paragraph(f'Дата: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph(f'Найдено ошибок: {len(errors)}')
    doc.add_paragraph('')

    # Легенда
    legend = doc.add_paragraph()
    legend.add_run('Легенда: ').bold = True
    legend.add_run('КАПС — что услышал Яндекс, ')
    run = legend.add_run('жирный')
    run.bold = True
    legend.add_run(' — правильное слово')

    doc.add_paragraph('_' * 50)

    # Ошибки
    for error in errors:
        time = error.get('time', 0)
        mins = int(time // 60)
        secs = int(time % 60)
        time_str = f"{mins}:{secs:02d}"

        error_type = error.get('type', '')
        original = error.get('original', '')
        transcript = error.get('transcript', '')
        context = error.get('context', '')

        p = doc.add_paragraph()

        # Время
        run = p.add_run(f"{time_str} — ")
        run.bold = True

        # Тип ошибки
        if error_type == 'substitution':
            run = p.add_run(transcript.upper())
            run.italic = True
            p.add_run(' → ')
            run = p.add_run(original)
            run.bold = True
        elif error_type == 'deletion':
            p.add_run('Пропущено: ')
            run = p.add_run(f'({original})')
            run.bold = True
        elif error_type == 'insertion':
            p.add_run('Лишнее: ')
            run = p.add_run(transcript.upper())
            run.italic = True

        # Контекст
        if context:
            p = doc.add_paragraph()
            p.add_run(f'   ...{context[:100]}...')
            p.paragraph_format.left_indent = Inches(0.5)

        doc.add_paragraph('')

    # Сохраняем с использованием FileNaming
    report_path = Path(report_path)
    if HAS_CONFIG:
        chapter_id = FileNaming.get_chapter_id(report_path)
        output_path = Path(output_dir) / FileNaming.build_filename(chapter_id, 'docx')
    else:
        output_path = Path(output_dir) / (report_path.stem.replace('_filtered', '').replace('_final', '') + '_для_чтеца.docx')

    doc.save(str(output_path))
    print(f"  DOCX сохранён: {output_path}")

    return str(output_path)


def run_pipeline(audio_path, text_path, output_dir=None,
                 skip_convert=False, skip_transcribe=False,
                 transcript_path=None, threshold=None, web=False,
                 show_progress=True, phantom=None, force=False):
    """
    Запускает полный пайплайн.

    Args:
        audio_path: путь к аудиофайлу
        text_path: путь к оригинальному тексту
        output_dir: папка для результатов
        skip_convert: пропустить конвертацию
        skip_transcribe: пропустить транскрибацию (если есть JSON)
        transcript_path: путь к готовой транскрипции
        threshold: порог схожести (None = из SmartCompareConfig)
        web: открыть результат в браузере
        show_progress: показывать прогресс-бар (tqdm)
        phantom: пропустить первые N секунд (None = из SmartCompareConfig)
        force: перезаписать существующие результаты
    """
    global _progress

    audio_path = Path(audio_path)
    text_path = Path(text_path)

    # ---- Валидация входных данных ----
    if not audio_path.exists():
        raise FileNotFoundError(f"Аудиофайл не найден: {audio_path}")
    if not text_path.exists():
        raise FileNotFoundError(f"Текстовый файл не найден: {text_path}")

    # Проверка: аудиофайл не пуст
    if audio_path.stat().st_size == 0:
        raise ValueError(f"Аудиофайл пуст (0 байт): {audio_path}")

    # Проверка: текстовый файл не пуст
    if text_path.stat().st_size == 0:
        raise ValueError(f"Текстовый файл пуст (0 байт): {text_path}")

    # Проверка: допустимое расширение аудио
    valid_audio_ext = {'.mp3', '.ogg', '.opus', '.wav', '.raw', '.pcm'}
    if audio_path.suffix.lower() not in valid_audio_ext:
        raise ValueError(f"Неподдерживаемый формат аудио: {audio_path.suffix}. "
                         f"Допустимые: {', '.join(sorted(valid_audio_ext))}")

    # Проверка: допустимое расширение текста
    valid_text_ext = {'.docx', '.txt'}
    if text_path.suffix.lower() not in valid_text_ext:
        raise ValueError(f"Неподдерживаемый формат текста: {text_path.suffix}. "
                         f"Допустимые: {', '.join(sorted(valid_text_ext))}")

    # Проверка: DOCX не повреждён (минимальная валидация)
    if text_path.suffix.lower() == '.docx':
        try:
            from docx import Document
            doc = Document(str(text_path))
            text_content = ' '.join(p.text for p in doc.paragraphs).strip()
            if not text_content:
                raise ValueError(f"DOCX файл не содержит текста: {text_path}")
        except Exception as e:
            if 'не содержит' in str(e):
                raise
            raise ValueError(f"DOCX файл повреждён или не читается: {text_path}: {e}")

    # Логируем начало пайплайна
    logger.info(f"{'='*60}")
    logger.info(f"Запуск пайплайна v{VERSION}")
    logger.info(f"Аудио: {audio_path}")
    logger.info(f"Текст: {text_path}")
    logger.debug(f"Параметры: skip_convert={skip_convert}, skip_transcribe={skip_transcribe}, "
                 f"threshold={threshold}, phantom={phantom}, force={force}")

    # Папка для результатов (используем RESULTS_DIR из config.py)
    if output_dir:
        output_dir = Path(output_dir)
    else:
        output_dir = RESULTS_DIR / audio_path.stem

    # Проверяем существование папки результатов
    if output_dir.exists() and not force:
        action = 'overwrite' if force else 'ask'
        check_file_exists(output_dir, action=action)

    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n{'#'*60}")
    print(f"  ПАЙПЛАЙН ПРОВЕРКИ АУДИОКНИГИ")
    print(f"{'#'*60}")
    print(f"  Аудио: {audio_path}")
    print(f"  Текст: {text_path}")
    print(f"  Результаты: {output_dir}")
    print(f"{'#'*60}")

    # Создаём прогресс-бар
    skip_steps = []
    if skip_convert or skip_transcribe:
        skip_steps.append("Конвертация в OggOpus")
    if skip_transcribe:
        skip_steps.append("Транскрибация (Яндекс SpeechKit)")

    _progress = PipelineProgress(
        use_tqdm=show_progress and HAS_TQDM,
        skip_steps=skip_steps
    )

    results = {}

    # Шаг 1: Конвертация
    if not skip_convert and not skip_transcribe:
        converted_audio = run_step(
            "Конвертация в OggOpus",
            step_convert_audio,
            str(audio_path), str(output_dir)
        )
        results['converted_audio'] = converted_audio
    else:
        converted_audio = str(audio_path)
        results['converted_audio'] = converted_audio

    # Шаг 2: Транскрибация
    if skip_transcribe and transcript_path:
        # Валидируем готовую транскрипцию
        print(f"\n  Валидация транскрипции: {transcript_path}")
        try:
            validate_transcript_file(transcript_path)
        except TranscriptValidationError as e:
            print(f"  ✗ Ошибка валидации: {e}")
            if e.details.get('errors'):
                for err in e.details['errors'][:5]:
                    print(f"    - {err}")
            raise
        results['transcript'] = transcript_path
        print(f"  Используем готовую транскрипцию: {transcript_path}")
    elif not skip_transcribe:
        results['transcript'] = run_step(
            "Транскрибация (Яндекс SpeechKit)",
            step_transcribe,
            converted_audio, str(output_dir)
        )
    else:
        # Ищем JSON рядом с аудио
        json_path = audio_path.with_suffix('.json')
        if json_path.exists():
            results['transcript'] = str(json_path)
        else:
            raise FileNotFoundError(
                "Транскрипция не найдена. Укажите --transcript-path или уберите --skip-transcribe"
            )

    # Шаг 3.5: Проверка соответствия транскрипции и текста
    results['correspondence_valid'] = run_step(
        "Проверка соответствия файлов",
        step_validate_correspondence,
        results['transcript'], str(text_path)
    )

    # Шаг 3: Нормализация текста
    results['normalized_text'] = run_step(
        "Нормализация текста",
        step_normalize_text,
        str(text_path), str(output_dir)
    )

    # Шаг 4: Умное сравнение
    # Используем оригинальный DOCX для контекстов с пунктуацией
    results['compared'] = run_step(
        "Умное сравнение (якоря + серые зоны)",
        step_smart_compare,
        results['transcript'], str(text_path),
        str(audio_path), str(output_dir), threshold, phantom
    )

    # Шаг 5: Фильтр отсева
    results['filtered'] = run_step(
        "Фильтр отсева",
        step_golden_filter,
        results['compared'], str(output_dir)
    )

    # Шаг 6: DOCX для чтеца
    results['docx'] = run_step(
        "Генерация DOCX для чтеца",
        step_generate_docx,
        results['filtered'], str(audio_path), str(output_dir)
    )

    # Закрываем прогресс-бар
    if _progress:
        _progress.close()

    # Итоги
    print(f"\n{'#'*60}")
    print(f"  ГОТОВО!")
    print(f"{'#'*60}")
    print(f"  Результаты в: {output_dir}")

    # Загружаем финальный отчёт для статистики
    with open(results['filtered'], 'r', encoding='utf-8') as f:
        final_report = json.load(f)

    errors = final_report.get('errors', [])
    print(f"  Найдено ошибок чтеца: {len(errors)}")

    # Логируем итоги
    logger.info(f"{'='*60}")
    logger.info(f"Пайплайн завершён успешно")
    logger.info(f"Результаты: {output_dir}")
    logger.info(f"Найдено ошибок: {len(errors)}")

    if results['docx']:
        print(f"  DOCX для чтеца: {results['docx']}")

    # Открываем в браузере
    if web:
        print(f"\n  Открываю веб-интерфейс...")
        try:
            from web_viewer import main as web_main
            # Запускаем веб-сервер с оригиналом и транскрипцией для обогащения контекстов
            cmd = [
                sys.executable,
                str(SCRIPT_DIR / 'web_viewer.py'),
                results['filtered'],
                '--audio', str(audio_path),
                '--original', str(text_path),
                '--transcript', results['transcript']
            ]
            subprocess.Popen(cmd)
        except Exception as e:
            print(f"  ⚠ Не удалось открыть веб-интерфейс: {e}")

    return results


def main():
    global logger

    parser = argparse.ArgumentParser(
        description='Полный пайплайн проверки аудиокниги',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=f"""
Этапы:
  1. Конвертация аудио → OggOpus моно
  2. Транскрибация → Яндекс SpeechKit
  3. Нормализация текста → числа, сокращения, дефисы
  4. Умное сравнение → якоря + серые зоны
  5. Фильтр отсева → Левенштейн, лемматизация
  6. Генерация DOCX → отчёт для чтеца

Примеры:
  python pipeline.py глава.mp3 оригинал.docx
  python pipeline.py глава.mp3 оригинал.txt --skip-transcribe --transcript-path транскрипт.json
  python pipeline.py глава.mp3 оригинал.docx --web --force
  python pipeline.py глава.mp3 оригинал.docx --verbose  # подробный вывод

Текущие настройки из config.py:
  threshold: {SmartCompareConfig.THRESHOLD if HAS_CONFIG else 0.7}
  phantom_seconds: {SmartCompareConfig.PHANTOM_SECONDS if HAS_CONFIG else -1} (-1 = авто)
  results_dir: {RESULTS_DIR}
        """
    )
    parser.add_argument('audio', nargs='?', help='Аудиофайл (MP3, OGG, WAV)')
    parser.add_argument('text', nargs='?', help='Оригинальный текст (TXT или DOCX)')
    parser.add_argument('--output', '-o', help='Папка для результатов')
    parser.add_argument('--skip-convert', action='store_true',
                        help='Пропустить конвертацию аудио')
    parser.add_argument('--skip-transcribe', action='store_true',
                        help='Пропустить транскрибацию')
    parser.add_argument('--transcript-path', '-t',
                        help='Путь к готовой транскрипции (JSON)')
    parser.add_argument('--threshold', type=float, default=None,
                        help=f'Порог схожести 0-1 (по умолчанию: {SmartCompareConfig.THRESHOLD if HAS_CONFIG else 0.7})')
    parser.add_argument('--web', '-w', action='store_true',
                        help='Открыть результат в веб-интерфейсе')
    parser.add_argument('--no-progress', action='store_true',
                        help='Отключить прогресс-бар')
    parser.add_argument('--phantom', type=float, default=None,
                        help='Пропустить первые N секунд транскрипции (метаданные). -1 = автоопределение')
    parser.add_argument('--force', '-f', action='store_true',
                        help='Перезаписать существующие результаты')
    parser.add_argument('--verbose', '-v', action='store_true',
                        help='Подробный вывод (уровень DEBUG)')
    parser.add_argument('--quiet', '-q', action='store_true',
                        help='Только ошибки (уровень ERROR)')
    parser.add_argument('--version', '-V', action='store_true',
                        help='Показать версию и выйти')

    args = parser.parse_args()

    # Показать версию и выйти
    if args.version:
        print(f"Pipeline v{VERSION} ({VERSION_DATE})")
        print(f"  config.py: {'подключён' if HAS_CONFIG else 'fallback'}")
        print(f"  logging: {'из config.py' if HAS_LOGGING else 'fallback'}")
        return

    # Настраиваем уровень логирования
    if args.quiet:
        log_level = 'ERROR'
    elif args.verbose:
        log_level = 'DEBUG'
    else:
        log_level = 'INFO'

    # Инициализируем систему логирования
    logger = setup_logging(
        level=log_level,
        module_name='pipeline',
        console=True
    )

    # Очистка старых логов
    if HAS_LOGGING:
        try:
            from config import LOGS_DIR
            removed = cleanup_old_logs(LOGS_DIR)
            if removed > 0:
                logger.debug(f"Удалено старых логов: {removed}")
        except Exception:
            pass

    # Проверяем обязательные аргументы
    if not args.audio or not args.text:
        parser.print_help()
        print("\n✗ Ошибка: требуются аргументы audio и text")
        sys.exit(1)

    print(f"\n{'#'*60}")
    print(f"  Pipeline v{VERSION}")
    print(f"{'#'*60}")

    try:
        run_pipeline(
            args.audio,
            args.text,
            output_dir=args.output,
            skip_convert=args.skip_convert,
            skip_transcribe=args.skip_transcribe,
            transcript_path=args.transcript_path,
            threshold=args.threshold,
            web=args.web,
            show_progress=not args.no_progress,
            phantom=args.phantom,
            force=args.force
        )
    except FileNotFoundError as e:
        logger.error(f"Файл не найден: {e}")
        print(f"\n✗ Файл не найден: {e}")
        sys.exit(1)
    except TranscriptValidationError as e:
        logger.error(f"Ошибка валидации транскрипции: {e}")
        if e.details.get('errors'):
            for err in e.details['errors'][:5]:
                logger.error(f"  - {err}")
        print(f"\n✗ Ошибка валидации: {e}")
        sys.exit(1)
    except KeyboardInterrupt:
        logger.warning("Прервано пользователем (Ctrl+C)")
        print(f"\n⚠ Прервано пользователем")
        sys.exit(130)
    except Exception as e:
        log_exception(logger, "Критическая ошибка пайплайна", e)
        print(f"\n✗ Ошибка: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        else:
            print("  Используйте --verbose для детальной информации")
        sys.exit(1)


if __name__ == '__main__':
    main()
