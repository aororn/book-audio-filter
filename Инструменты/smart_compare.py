#!/usr/bin/env python3
"""
Smart Compare v10.0 - Умное сравнение транскрипции с оригиналом

Алгоритм:
1. МАКРО-ВЫРАВНИВАНИЕ: Находим якоря (уникальные слова >6 символов)
2. СЕГМЕНТАЦИЯ: Делим текст на сегменты по якорям
3. МИКРО-ВЫРАВНИВАНИЕ: SequenceMatcher на каждый сегмент отдельно
4. ОБЪЕДИНЕНИЕ: Собираем opcodes с корректировкой индексов
5. КЛАССИФИКАЦИЯ: Отделяем ошибки Яндекса от ошибок чтеца

Критерий различия:
- Малое расстояние Левенштейна → фонетическая ошибка Яндекса (игнорируем)
- Большое расстояние → реальная ошибка чтеца (в отчёт)

Дополнительно: фонетическое сравнение для русского языка

Использование:
    python smart_compare.py транскрипция.json оригинал.txt --audio глава.mp3
    python smart_compare.py транскрипция.json оригинал.txt --threshold 0.7
    python smart_compare.py транскрипция.json оригинал.txt --force

Changelog:
    v10.1 (2026-01-29): Интеграция WindowVerifier
        - Добавлен WindowVerifier для контекстной верификации ошибок
        - Sliding Window: проверка контекста без учёта пробелов
        - Если контекст ≥95% идентичен — это технический шум
    v10.0 (2026-01-29): Полное посегментное выравнивание
        - compare_with_anchors() — посегментное сравнение с явным включением якорей
        - merge_adjacent_opcodes() — объединение смежных opcodes
        - Якоря теперь включаются в opcodes как 'equal'
        - Полный переход от полнотекстового к макро/микро выравниванию
    v9.1 (2026-01-29): Fix ComparisonStitcher — защита коротких слов
        - PROTECTED_SHORT_WORDS: частицы, предлоги, союзы не склеиваются
        - Исправлена регрессия: "кто"+"то" больше не склеивается в "ктото"
        - Golden тест: восстановлено 93/93 (100%)
    v9.0 (2026-01-29): Интеграция AlignmentManager и ScoringEngine
        - Посегментное выравнивание через AlignmentManager
        - Адаптивные штрафы через ScoringEngine
        - РЕГРЕССИЯ: потеря 3 golden ошибок из-за Stitcher
    v7.0 (2026-01-29): ComparisonStitcher — склейка разбитых слов
        - Добавлен ComparisonStitcher для склейки слов типа "средо"+"точие"
        - Интеграция с CharacterGuard для увеличенного буфера имён
        - Нормализация дефисов: "Красно-волосый" → "Красноволосый"
    v6.0 (2026-01-26): Интеграция с smart_rules и улучшенная фонетика
        - Используем phonetic_normalize из filters.smart_rules
        - Улучшенный анализ серых зон
        - Интеграция с morphology.py v6.0
    v3.1 (2026-01-24): Исправление неправильных сопоставлений
        - fix_misaligned_errors(): пост-обработка ошибок SequenceMatcher
        - Исправляет случаи типа "рагидон"→"и" + deletion "рагедон"
    v3.0 (2026-01-24): Интеграция с config.py
        - SmartCompareConfig для threshold/phantom_seconds
        - FileNaming для выходных файлов
        - check_file_exists() + флаг --force
        - Интеграция с morphology.py для лемматизации
    v2.0: Базовая версия с SequenceMatcher
"""

# Версия модуля
VERSION = '10.1.0'
VERSION_DATE = '2026-01-29'

import argparse
import json
import re
import os
from pathlib import Path
from dataclasses import dataclass, asdict
from typing import List, Tuple, Optional
from difflib import SequenceMatcher


# =============================================================================
# ИМПОРТ ЦЕНТРАЛИЗОВАННОЙ КОНФИГУРАЦИИ
# =============================================================================

try:
    from config import (
        FileNaming, SmartCompareConfig, check_file_exists
    )
    HAS_CONFIG = True
except ImportError:
    HAS_CONFIG = False

# Импорт морфологии из централизованного модуля
try:
    from morphology import get_lemma as morphology_get_lemma, HAS_PYMORPHY
except ImportError:
    morphology_get_lemma = None
    HAS_PYMORPHY = False

# Импорт smart_rules (v6.0)
try:
    from filters.smart_rules import (
        get_smart_rules, is_smart_false_positive, phonetic_normalize as smart_phonetic_normalize
    )
    HAS_SMART_RULES = True
except ImportError:
    HAS_SMART_RULES = False
    smart_phonetic_normalize = None


# =============================================================================
# ИМПОРТ ЗАВИСИМОСТЕЙ
# =============================================================================

try:
    from thefuzz import fuzz
    HAS_THEFUZZ = True
except ImportError:
    try:
        from fuzzywuzzy import fuzz
        HAS_THEFUZZ = True
    except ImportError:
        HAS_THEFUZZ = False
        print("⚠ Установите: pip install thefuzz")

try:
    from rapidfuzz.distance import Levenshtein
    HAS_RAPIDFUZZ = True
except ImportError:
    HAS_RAPIDFUZZ = False

# Импорт CharacterGuard для защиты имён (v9.0)
try:
    from filters.character_guard import get_character_guard, is_character_name
    HAS_CHARACTER_GUARD = True
except ImportError:
    HAS_CHARACTER_GUARD = False
    is_character_name = lambda x: False

# Импорт AlignmentManager для сегментации (v9.0)
try:
    from alignment_manager import AlignmentManager, segment_texts, AnchorPoint, Segment
    HAS_ALIGNMENT_MANAGER = True
except ImportError:
    HAS_ALIGNMENT_MANAGER = False
    AnchorPoint = None
    Segment = None

# Импорт ScoringEngine для адаптивных штрафов (v9.0)
try:
    from filters.scoring_engine import (
        get_scoring_engine, should_filter_by_score, calculate_penalty
    )
    HAS_SCORING_ENGINE = True
except ImportError:
    HAS_SCORING_ENGINE = False

# Импорт WindowVerifier для верификации сегментов (v10.1)
try:
    from filters.window_verifier import (
        get_window_verifier, verify_segment, is_technical_noise,
        VerificationStatus, VerificationResult
    )
    HAS_WINDOW_VERIFIER = True
except ImportError:
    HAS_WINDOW_VERIFIER = False


# =============================================================================
# ПОСЕГМЕНТНОЕ ВЫРАВНИВАНИЕ v10.0
# =============================================================================

def merge_adjacent_opcodes(opcodes: list) -> list:
    """
    Объединяет соседние opcodes с одинаковым тегом.

    Это необходимо после посегментного сравнения, чтобы
    объединить 'equal' блоки из разных сегментов.

    Args:
        opcodes: список кортежей (tag, i1, i2, j1, j2)

    Returns:
        объединённый список opcodes
    """
    if not opcodes:
        return opcodes

    merged = [list(opcodes[0])]
    for tag, i1, i2, j1, j2 in opcodes[1:]:
        prev = merged[-1]
        # Если тот же тег и позиции смежные
        if tag == prev[0] and i1 == prev[2] and j1 == prev[4]:
            merged[-1] = [tag, prev[1], i2, prev[3], j2]
        else:
            merged.append([tag, i1, i2, j1, j2])

    return [tuple(op) for op in merged]


def compare_with_anchors(
    orig_norm: list,
    trans_norm: list,
    anchors: list,
    segments: list,
    original_words: list = None,
    transcript_words: list = None
) -> tuple:
    """
    Посегментное сравнение с явным включением якорей.

    Алгоритм:
    1. Для каждого сегмента выполняем SequenceMatcher
    2. Корректируем индексы opcodes на глобальные позиции
    3. Добавляем якорь после каждого сегмента как 'equal' opcode
    4. Объединяем смежные opcodes с одинаковым тегом

    Args:
        orig_norm: нормализованные слова оригинала
        trans_norm: нормализованные слова транскрипции
        anchors: список AnchorPoint
        segments: список Segment
        original_words: Word объекты оригинала (для отладки)
        transcript_words: Word объекты транскрипции (для отладки)

    Returns:
        (opcodes, ratio) — список opcodes и коэффициент схожести
    """
    all_opcodes = []
    total_matches = 0
    total_len = len(orig_norm) + len(trans_norm)

    # Создаём словарь якорей по слову для быстрого доступа
    anchor_dict = {a.word: a for a in anchors}

    for seg in segments:
        # Извлекаем слова сегмента
        seg_orig = orig_norm[seg.orig_start:seg.orig_end]
        seg_trans = trans_norm[seg.trans_start:seg.trans_end]

        if seg_orig or seg_trans:
            # Сравниваем сегмент
            matcher = SequenceMatcher(None, seg_orig, seg_trans)
            seg_opcodes = matcher.get_opcodes()

            # Корректируем индексы на глобальные позиции
            for tag, i1, i2, j1, j2 in seg_opcodes:
                global_opcode = (
                    tag,
                    i1 + seg.orig_start,
                    i2 + seg.orig_start,
                    j1 + seg.trans_start,
                    j2 + seg.trans_start
                )
                all_opcodes.append(global_opcode)

                if tag == 'equal':
                    total_matches += (i2 - i1) + (j2 - j1)

        # КЛЮЧЕВОЕ: Добавляем якорь ПОСЛЕ сегмента как 'equal'
        # Это исправляет проблему, когда якоря исключались из сравнения
        if seg.anchor_after and seg.anchor_after in anchor_dict:
            anchor = anchor_dict[seg.anchor_after]
            anchor_opcode = (
                'equal',
                anchor.orig_idx,
                anchor.orig_idx + 1,
                anchor.trans_idx,
                anchor.trans_idx + 1
            )
            all_opcodes.append(anchor_opcode)
            total_matches += 2  # Слово в orig + слово в trans

    # Сортируем opcodes по позиции в оригинале (и в транскрипции для равных)
    all_opcodes.sort(key=lambda x: (x[1], x[3]))

    # Объединяем соседние opcodes с одинаковым тегом
    all_opcodes = merge_adjacent_opcodes(all_opcodes)

    # Вычисляем ratio
    ratio = total_matches / total_len if total_len > 0 else 1.0

    return all_opcodes, ratio


# =============================================================================
# COMPARISON STITCHER v1.0 — Склейка разбитых слов (v9.0)
# =============================================================================

class ComparisonStitcher:
    """
    Склеивает слова, которые Яндекс разбил на части.

    Примеры:
    - "средо" + "точие" → "средоточие"
    - "Красно" + "волосый" → "Красноволосый"
    - "Ми" + "ражный" → "Миражный"

    Использует CharacterGuard для увеличения буфера при склейке имён.

    v1.1 (2026-01-29): Добавлена защита коротких служебных слов от склейки
    """

    # Минимальная длина фрагмента для склейки
    MIN_FRAGMENT_LEN = 2

    # Максимальный буфер вперёд (сколько слов проверять)
    DEFAULT_BUFFER = 2
    NAME_BUFFER = 4  # увеличенный буфер для имён персонажей

    # Порог схожести для склейки (Левенштейн ratio)
    STITCH_THRESHOLD = 0.85

    # Защищённые короткие слова — НЕ склеивать их с соседними
    # Это частицы, предлоги, союзы, которые могут быть insertion ошибками
    PROTECTED_SHORT_WORDS = frozenset({
        'то', 'по', 'на', 'за', 'от', 'до', 'ни', 'не', 'бы', 'же', 'ли',
        'и', 'а', 'я', 'у', 'о', 'в', 'к', 'с',
        'ещё', 'еще', 'уже', 'вот', 'вон', 'тут', 'там',
    })

    def __init__(self, original_words: list = None):
        """
        Args:
            original_words: список слов из оригинала для поиска целевых слов
        """
        self._original_set: set = set()
        self._original_normalized: set = set()
        if original_words:
            for w in original_words:
                text = w.text if hasattr(w, 'text') else str(w)
                norm = w.normalized if hasattr(w, 'normalized') else text.lower()
                self._original_set.add(text.lower())
                self._original_normalized.add(norm)

    def pre_stitch_audio_words(self, words: list) -> list:
        """
        Основной метод: склеивает разбитые слова в списке.

        Args:
            words: список Word объектов из транскрипции

        Returns:
            Новый список с склеенными словами
        """
        if not words:
            return words

        result = []
        i = 0

        while i < len(words):
            word = words[i]
            text = word.text if hasattr(word, 'text') else str(word)
            norm = word.normalized if hasattr(word, 'normalized') else text.lower()

            # Проверяем, не начало ли это разбитого слова
            buffer = self._get_buffer(text)
            merged = self._try_merge(words, i, buffer)

            if merged:
                # Создаём новый Word объект с склеенным текстом
                merged_word = self._create_merged_word(words, i, merged)
                result.append(merged_word)
                i += merged['count']
            else:
                result.append(word)
                i += 1

        return result

    def _get_buffer(self, text: str) -> int:
        """Определяет размер буфера для проверки склейки."""
        if HAS_CHARACTER_GUARD and is_character_name(text):
            return self.NAME_BUFFER
        return self.DEFAULT_BUFFER

    def _try_merge(self, words: list, start_idx: int, buffer: int) -> dict | None:
        """
        Пытается склеить слово с последующими.

        Returns:
            dict с 'text', 'normalized', 'count' или None
        """
        if start_idx >= len(words):
            return None

        base_word = words[start_idx]
        base_text = base_word.text if hasattr(base_word, 'text') else str(base_word)

        # Слишком короткий фрагмент — не склеиваем
        if len(base_text) < self.MIN_FRAGMENT_LEN:
            return None

        # ЗАЩИТА: Базовое слово тоже может быть защищённым
        # Например, "кто" не должно склеиваться с "то" в "ктото"
        if base_text.lower() in self.PROTECTED_SHORT_WORDS:
            return None

        # Пробуем склеить с 1, 2, ... buffer следующих слов
        for count in range(1, min(buffer + 1, len(words) - start_idx)):
            merged_parts = [base_text]
            should_skip = False

            for j in range(1, count + 1):
                next_word = words[start_idx + j]
                next_text = next_word.text if hasattr(next_word, 'text') else str(next_word)

                # ЗАЩИТА: Не склеивать защищённые короткие слова
                # Они могут быть insertion ошибками (лишние слова)
                if next_text.lower() in self.PROTECTED_SHORT_WORDS:
                    should_skip = True
                    break

                merged_parts.append(next_text)

            # Если встретили защищённое слово — пропускаем эту попытку склейки
            if should_skip:
                continue

            merged_text = ''.join(merged_parts)
            merged_norm = merged_text.lower().replace('ё', 'е')

            # Нормализация дефисов
            merged_no_hyphen = merged_text.replace('-', '')
            merged_norm_no_hyphen = merged_no_hyphen.lower().replace('ё', 'е')

            # Проверяем: есть ли такое слово в оригинале?
            if merged_norm in self._original_normalized:
                return {
                    'text': merged_text,
                    'normalized': merged_norm,
                    'count': count + 1
                }

            # Проверяем без дефиса (Красно-волосый → Красноволосый)
            if merged_norm_no_hyphen in self._original_normalized:
                return {
                    'text': merged_no_hyphen,
                    'normalized': merged_norm_no_hyphen,
                    'count': count + 1
                }

            # Проверяем через CharacterGuard (имена могут не быть в оригинале точно)
            if HAS_CHARACTER_GUARD and is_character_name(merged_text):
                return {
                    'text': merged_text,
                    'normalized': merged_norm,
                    'count': count + 1
                }

        return None

    def _create_merged_word(self, words: list, start_idx: int, merged: dict):
        """Создаёт новый Word объект из склеенных слов."""
        from dataclasses import replace

        first_word = words[start_idx]
        last_word = words[start_idx + merged['count'] - 1]

        # Если Word — dataclass, используем replace
        if hasattr(first_word, '__dataclass_fields__'):
            return replace(
                first_word,
                text=merged['text'],
                normalized=merged['normalized'],
                time_end=last_word.time_end if hasattr(last_word, 'time_end') else first_word.time_end
            )
        else:
            # Fallback: создаём новый объект
            return type(first_word)(
                text=merged['text'],
                normalized=merged['normalized'],
                position=first_word.position,
                time_start=first_word.time_start,
                time_end=last_word.time_end if hasattr(last_word, 'time_end') else first_word.time_end,
                confidence=first_word.confidence
            )

    def normalize_hyphens(self, words: list) -> list:
        """
        Нормализует дефисы в словах.

        "Красно-волосый" → "Красноволосый" (если в оригинале без дефиса)
        """
        result = []
        for word in words:
            text = word.text if hasattr(word, 'text') else str(word)

            if '-' in text:
                no_hyphen = text.replace('-', '')
                no_hyphen_norm = no_hyphen.lower().replace('ё', 'е')

                if no_hyphen_norm in self._original_normalized:
                    # Заменяем на версию без дефиса
                    if hasattr(word, '__dataclass_fields__'):
                        from dataclasses import replace
                        word = replace(word, text=no_hyphen, normalized=no_hyphen_norm)
                    else:
                        word.text = no_hyphen
                        word.normalized = no_hyphen_norm

            result.append(word)
        return result


# Singleton
_stitcher_instance = None

def get_stitcher(original_words: list = None) -> ComparisonStitcher:
    """Возвращает экземпляр Stitcher."""
    global _stitcher_instance
    if _stitcher_instance is None or original_words is not None:
        _stitcher_instance = ComparisonStitcher(original_words)
    return _stitcher_instance


def stitch_audio_words(audio_words: list, original_words: list = None) -> list:
    """Быстрая функция для склейки слов."""
    stitcher = get_stitcher(original_words)
    return stitcher.pre_stitch_audio_words(audio_words)


# =============================================================================
# ФОНЕТИКА РУССКОГО ЯЗЫКА
# =============================================================================

# Фонетические группы (звуки, которые часто путаются)
PHONETIC_GROUPS = {
    # Гласные
    'а': 'a', 'о': 'a',  # безударные о→а
    'е': 'i', 'и': 'i', 'я': 'i',  # безударные
    'э': 'e',
    'у': 'u', 'ю': 'u',
    'ы': 'y',

    # Звонкие/глухие согласные
    'б': 'p', 'п': 'p',
    'в': 'f', 'ф': 'f',
    'г': 'k', 'к': 'k', 'х': 'k',
    'д': 't', 'т': 't',
    'ж': 'sh', 'ш': 'sh', 'щ': 'sh',
    'з': 's', 'с': 's', 'ц': 's',
    'ч': 'ch',

    # Сонорные
    'л': 'l',
    'м': 'm',
    'н': 'n',
    'р': 'r',

    # Другие
    'й': 'j',
}


def to_phonetic(word: str) -> str:
    """
    Преобразует слово в фонетическое представление.
    Убирает различия, которые не слышны на слух.
    """
    word = word.lower().replace('ё', 'е')

    # Убираем мягкий и твёрдый знаки
    word = word.replace('ь', '').replace('ъ', '')

    # Преобразуем в фонетические группы
    result = []
    for char in word:
        if char in PHONETIC_GROUPS:
            result.append(PHONETIC_GROUPS[char])
        elif char.isalpha():
            result.append(char)

    # Убираем дубли (удвоенные согласные звучат как одна)
    cleaned = []
    for char in result:
        if not cleaned or cleaned[-1] != char:
            cleaned.append(char)

    return ''.join(cleaned)


def phonetic_similarity(word1: str, word2: str) -> float:
    """
    Вычисляет фонетическую схожесть двух слов (0-100).

    v6.0: Использует smart_phonetic_normalize если доступен.
    """
    # Используем улучшенную фонетику из smart_rules если доступна
    if HAS_SMART_RULES and smart_phonetic_normalize:
        p1 = smart_phonetic_normalize(word1)
        p2 = smart_phonetic_normalize(word2)
    else:
        p1 = to_phonetic(word1)
        p2 = to_phonetic(word2)

    if p1 == p2:
        return 100.0

    if HAS_THEFUZZ:
        return fuzz.ratio(p1, p2)
    elif HAS_RAPIDFUZZ:
        max_len = max(len(p1), len(p2))
        if max_len == 0:
            return 100.0
        dist = Levenshtein.distance(p1, p2)
        return (1 - dist / max_len) * 100
    else:
        # Простое сравнение
        matches = sum(1 for a, b in zip(p1, p2) if a == b)
        max_len = max(len(p1), len(p2))
        return (matches / max_len) * 100 if max_len > 0 else 100.0


# =============================================================================
# НОРМАЛИЗАЦИЯ — импорт из единого источника morphology.py
# =============================================================================

# Импорт normalize_word из morphology.py
try:
    from morphology import normalize_word as _normalize_word_base
    HAS_MORPHOLOGY = True
except ImportError:
    HAS_MORPHOLOGY = False
    def _normalize_word_base(word: str) -> str:
        """Fallback нормализация."""
        return word.lower().strip().replace('ё', 'е')


def normalize_word(word: str) -> str:
    """
    Нормализует слово для сравнения.
    Использует morphology.normalize_word + убирает пунктуацию.
    """
    word = _normalize_word_base(word)
    # Дополнительно убираем пунктуацию для сравнения
    word = re.sub(r'[^\w]', '', word)
    return word


def normalize_text(text: str) -> List[str]:
    """Разбивает текст на нормализованные слова"""
    # Убираем пунктуацию, приводим к нижнему регистру
    text = text.lower().replace('ё', 'е')
    text = re.sub(r'[^\w\s]', ' ', text)
    words = text.split()
    return [w for w in words if w]


def get_lemma(word: str) -> str:
    """Получает лемму слова (использует morphology.py если доступен)"""
    if morphology_get_lemma:
        return morphology_get_lemma(word)
    # Fallback: просто нормализация
    return normalize_word(word)


# =============================================================================
# СТРУКТУРЫ ДАННЫХ
# =============================================================================

@dataclass
class Word:
    """Слово с метаданными"""
    text: str
    normalized: str
    position: int
    time_start: float = 0.0
    time_end: float = 0.0
    confidence: float = 1.0
    original_text: str = ""  # оригинальный текст с пунктуацией (для контекста)


@dataclass
class Error:
    """Найденная ошибка"""
    type: str  # substitution, insertion, deletion
    time: float
    time_end: float = 0.0
    original: str = ""
    transcript: str = ""
    context: str = ""  # контекст из оригинала (основной)
    transcript_context: str = ""  # контекст из транскрипции (для insertions)
    marker_pos: int = -1  # позиция маркера в контексте (символы)
    confidence: float = 1.0
    is_yandex_error: bool = False
    similarity: float = 0.0
    phonetic_similarity: float = 0.0


# =============================================================================
# ПАРСИНГ ТРАНСКРИПЦИИ ЯНДЕКСА
# =============================================================================

def detect_phantom_seconds(transcript_json: str, original_text: str, max_phantom: float = 30.0) -> float:
    """
    Автоматически определяет длительность метаинформации в начале аудио.

    Сравнивает первые слова транскрипции с началом оригинального текста.
    Ищет позицию в транскрипции, где начинается настоящий текст книги.

    Алгоритм:
    1. Берём первые 6 слов оригинала
    2. Ищем в транскрипции позицию, где 4+ из 6 слов совпадают подряд
    3. Возвращаем время начала первого совпавшего слова

    Args:
        transcript_json: путь к JSON с транскрипцией
        original_text: путь к нормализованному тексту
        max_phantom: максимальная длительность фантома (сек)

    Returns:
        Количество секунд для пропуска (0 если не нужен пропуск)
    """
    import json
    from pathlib import Path

    # Загружаем транскрипцию
    with open(transcript_json, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Извлекаем слова с таймкодами
    trans_words = []
    for chunk in data.get('chunks', data.get('result', [])):
        for alt in chunk.get('alternatives', []):
            for w in alt.get('words', []):
                text = w.get('word', '').lower()
                start = w.get('startTime', '0s')
                if isinstance(start, str):
                    start = float(start.rstrip('s'))
                if text and start <= max_phantom:
                    trans_words.append((start, text))

    if not trans_words:
        return 0.0

    # Загружаем оригинал
    original_path = Path(original_text)
    if original_path.suffix.lower() == '.txt':
        with open(original_path, 'r', encoding='utf-8') as f:
            orig_text = f.read()
    else:
        # Для docx нужен python-docx
        try:
            from docx import Document
            doc = Document(str(original_path))
            orig_text = ' '.join(p.text for p in doc.paragraphs)
        except:
            return 0.0

    # Нормализуем первые слова оригинала (фильтруем пустые)
    orig_words = [normalize_word(w) for w in orig_text.split()[:50]]
    orig_words = [w for w in orig_words if w]

    if len(orig_words) < 5:
        return 0.0

    # Берём первые 6 слов оригинала для поиска
    check_words = orig_words[0:6]

    def words_similar(w1: str, w2: str) -> bool:
        """Проверяет похожесть двух слов."""
        if w1 == w2:
            return True
        # Для коротких слов требуем точное совпадение
        if len(w1) < 3 or len(w2) < 3:
            return False
        # Для длинных слов допускаем 1 ошибку
        if abs(len(w1) - len(w2)) <= 1:
            diff = sum(a != b for a, b in zip(w1, w2))
            # Учитываем разницу в длине
            diff += abs(len(w1) - len(w2))
            return diff <= 1
        return False

    # Ищем в транскрипции последовательность, похожую на начало оригинала
    for i, (time, word) in enumerate(trans_words):
        if time > max_phantom:
            break

        # Проверяем совпадение слов подряд, начиная с позиции i
        matches = 0
        for j, orig_word in enumerate(check_words):
            if i + j < len(trans_words):
                trans_word = trans_words[i + j][1]
                if words_similar(trans_word, orig_word):
                    matches += 1

        # Если совпало 4+ из 6 слов — нашли начало настоящего текста
        if matches >= 4:
            # Возвращаем время первого слова, если оно не в самом начале
            # (если первое слово текста на 0.5+ сек — значит есть вступление)
            if time < 0.5:
                return 0.0  # Текст начинается сразу, phantom не нужен
            return time

    # Не нашли совпадение — возвращаем 0
    return 0.0


def parse_yandex_transcription(json_path: str, phantom_seconds: float = 0) -> List[Word]:
    """
    Парсит JSON с транскрипцией от Яндекса.
    Извлекает слова с таймкодами.

    Args:
        json_path: путь к JSON файлу
        phantom_seconds: пропустить первые N секунд (метаданные/заставка)
    """
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    words = []
    position = 0

    # Формат ответа Яндекса: chunks → alternatives → words
    chunks = data.get('chunks', data.get('result', []))

    for chunk in chunks:
        alternatives = chunk.get('alternatives', [])
        if not alternatives:
            continue

        # Берём лучшую альтернативу
        best = alternatives[0]

        # Извлекаем слова с таймкодами
        chunk_words = best.get('words', [])

        for w in chunk_words:
            text = w.get('word', '')
            if not text:
                continue

            # Время в секундах (Яндекс возвращает в разных форматах)
            start_time = w.get('startTime', '0s')
            end_time = w.get('endTime', '0s')

            # Парсим время
            if isinstance(start_time, str):
                start_time = float(start_time.rstrip('s'))
            if isinstance(end_time, str):
                end_time = float(end_time.rstrip('s'))

            # Пропускаем фантом (метаданные в начале)
            if start_time < phantom_seconds:
                continue

            words.append(Word(
                text=text,
                normalized=normalize_word(text),
                position=position,
                time_start=start_time,
                time_end=end_time,
                confidence=best.get('confidence', 1.0)
            ))
            position += 1

    return words


def parse_original_text(text_path: str) -> List[Word]:
    """Парсит оригинальный текст, сохраняя оригинальные слова с пунктуацией"""
    # Читаем файл
    path = Path(text_path)
    if path.suffix.lower() == '.docx':
        try:
            from docx import Document
            doc = Document(str(path))
            text = '\n'.join(p.text for p in doc.paragraphs)
        except ImportError:
            raise ImportError("Установите: pip install python-docx")
    else:
        with open(path, 'r', encoding='utf-8') as f:
            text = f.read()

    # Извлекаем оригинальные слова с пунктуацией
    # Паттерн: слово + прилегающая пунктуация
    original_tokens = re.findall(r'[^\s]+', text)

    # Для каждого оригинального токена нормализуем и получаем список слов
    # Например: "что-то," → ["что", "то"]
    token_to_norm_words = []
    for token in original_tokens:
        # Нормализуем токен так же, как normalize_text
        token_lower = token.lower().replace('ё', 'е')
        token_clean = re.sub(r'[^\w\s]', ' ', token_lower)
        norm_words = [w for w in token_clean.split() if w]
        token_to_norm_words.append((token, norm_words))

    # Нормализуем для сравнения
    words_normalized = normalize_text(text)

    # Строим маппинг: для каждого нормализованного слова находим оригинальный токен
    words = []
    orig_idx = 0
    norm_word_idx_in_token = 0  # индекс внутри текущего составного токена

    for i, norm_word in enumerate(words_normalized):
        original_token = ""

        # Пропускаем токены без слов (только пунктуация)
        while orig_idx < len(token_to_norm_words):
            token, norm_words_in_token = token_to_norm_words[orig_idx]
            if not norm_words_in_token:
                orig_idx += 1
                norm_word_idx_in_token = 0
                continue
            break

        if orig_idx < len(token_to_norm_words):
            token, norm_words_in_token = token_to_norm_words[orig_idx]

            if norm_word_idx_in_token < len(norm_words_in_token):
                expected_norm = norm_words_in_token[norm_word_idx_in_token]

                if expected_norm == norm_word:
                    # Совпало — берём оригинальный токен
                    original_token = token
                    norm_word_idx_in_token += 1

                    # Если исчерпали все слова в токене, переходим к следующему
                    if norm_word_idx_in_token >= len(norm_words_in_token):
                        orig_idx += 1
                        norm_word_idx_in_token = 0
                else:
                    # Не совпало — возможно рассинхрон, пытаемся найти
                    found = False
                    for search_idx in range(orig_idx, min(orig_idx + 5, len(token_to_norm_words))):
                        search_token, search_norms = token_to_norm_words[search_idx]
                        if norm_word in search_norms:
                            original_token = search_token
                            word_pos_in_token = search_norms.index(norm_word)
                            # Переходим к этому токену
                            orig_idx = search_idx
                            norm_word_idx_in_token = word_pos_in_token + 1
                            if norm_word_idx_in_token >= len(search_norms):
                                orig_idx += 1
                                norm_word_idx_in_token = 0
                            found = True
                            break
                    if not found:
                        # Берём текущий токен как fallback
                        original_token = token

        words.append(Word(
            text=norm_word,
            normalized=normalize_word(norm_word),
            position=i,
            original_text=original_token
        ))

    return words


def get_context(words: List[Word], position: int, window: int = 10) -> Tuple[str, int]:
    """
    Получает контекст вокруг позиции с marker_pos НА слове.
    Для substitution и deletion — маркер указывает на само слово.

    Returns:
        (context_str, marker_pos) - контекст и позиция маркера в символах
    """
    start = max(0, position - window)
    end = min(len(words), position + window + 1)

    # Собираем слова до позиции
    before_words = []
    for w in words[start:position]:
        before_words.append(w.original_text if w.original_text else w.text)

    # Слово на позиции
    target_word = ""
    if position < len(words):
        w = words[position]
        target_word = w.original_text if w.original_text else w.text

    # Слова после позиции
    after_words = []
    for w in words[position + 1:end]:
        after_words.append(w.original_text if w.original_text else w.text)

    before_text = ' '.join(before_words)
    after_text = ' '.join(after_words)

    # Маркер ставится ПЕРЕД целевым словом
    marker_pos = len(before_text)
    if before_text:
        marker_pos += 1  # пробел

    # Собираем контекст
    parts = []
    if before_text:
        parts.append(before_text)
    if target_word:
        parts.append(target_word)
    if after_text:
        parts.append(after_text)

    context = ' '.join(parts)
    return context, marker_pos


def get_context_with_marker(words: List[Word], position: int, window: int = 10) -> Tuple[str, int]:
    """
    Получает контекст с позицией маркера (для insertion).

    Args:
        words: список слов оригинала
        position: позиция ДО которой должен быть вставлен маркер
        window: размер окна контекста

    Returns:
        (context_str, marker_pos) - контекст и позиция маркера в символах
    """
    start = max(0, position - window)
    end = min(len(words), position + window + 1)

    # Собираем слова до позиции и после
    before_words = []
    for w in words[start:position]:
        before_words.append(w.original_text if w.original_text else w.text)

    after_words = []
    for w in words[position:end]:
        after_words.append(w.original_text if w.original_text else w.text)

    before_text = ' '.join(before_words)
    after_text = ' '.join(after_words)

    # Маркер ставится после before_text (+ пробел если есть before)
    marker_pos = len(before_text)
    if before_text:
        marker_pos += 1  # пробел перед маркером

    context = before_text + (' ' if before_text and after_text else '') + after_text
    return context, marker_pos


def get_context_from_transcript(words: List[Word], position: int, window: int = 10) -> str:
    """Получает контекст из транскрипции"""
    start = max(0, position - window)
    end = min(len(words), position + window + 1)
    context_words = [w.text for w in words[start:end]]
    return ' '.join(context_words)


# =============================================================================
# ПОСТ-ОБРАБОТКА: ИСПРАВЛЕНИЕ НЕПРАВИЛЬНЫХ СОПОСТАВЛЕНИЙ
# =============================================================================

def fix_misaligned_errors(errors: List[Error]) -> List[Error]:
    """
    Исправляет неправильные сопоставления SequenceMatcher.

    Проблема: SequenceMatcher может сопоставить "рагидон" с "и", а потом
    создать deletion для "рагедон". Правильно: deletion "и", substitution "рагидон"→"рагедон".

    Паттерн для исправления:
    - substitution с низким сходством (< 0.4) рядом с deletion/insertion
    - слово из deletion/insertion похоже на слово из substitution
    """
    if not errors:
        return errors

    # Импортируем здесь чтобы избежать циклических импортов
    try:
        from rapidfuzz import fuzz
        has_fuzz = True
    except ImportError:
        has_fuzz = False
        return errors  # Без fuzz не можем проверять сходство

    fixed_errors = []
    skip_indices = set()

    # Сортируем по времени
    sorted_errors = sorted(enumerate(errors), key=lambda x: x[1].time)

    for i, (orig_idx, error) in enumerate(sorted_errors):
        if orig_idx in skip_indices:
            continue

        # Ищем substitution с низким сходством
        if error.type == 'substitution' and error.similarity < 0.4:
            trans_word = error.transcript.lower()
            orig_word = error.original.lower()

            # Ищем соседние deletion/insertion в пределах 2 секунд
            for j, (other_idx, other) in enumerate(sorted_errors):
                if other_idx in skip_indices or other_idx == orig_idx:
                    continue
                if abs(other.time - error.time) > 2.0:
                    continue

                if other.type == 'deletion':
                    del_word = other.original.lower()
                    # Проверяем: trans_word похоже на del_word?
                    sim = fuzz.ratio(trans_word, del_word) / 100.0

                    # Дополнительные проверки для избежания ложных пересопоставлений:
                    # 1. Не создаём substitution для одинаковых слов
                    # 2. Не меняем, если orig_word совпадает с trans_word
                    # 3. Оба слова должны быть достаточно длинными (>2 символов)
                    #    чтобы избежать ложных совпадений коротких слов
                    # 4. Слова должны быть похожей длины (разница <= 2)
                    len_diff = abs(len(trans_word) - len(del_word))
                    min_len = min(len(trans_word), len(del_word))

                    if (sim > 0.7 and
                        trans_word != orig_word and
                        del_word != orig_word and
                        min_len > 2 and
                        len_diff <= 2):
                        # Нашли! Это неправильное сопоставление.
                        # trans_word должно было сопоставиться с del_word
                        # А orig_word — это реальный пропуск (deletion)

                        # Создаём правильные ошибки:
                        # 1. substitution: trans_word → del_word (ошибка Яндекса/чтеца)
                        fixed_errors.append(Error(
                            type='substitution',
                            time=error.time,
                            time_end=error.time_end,
                            original=other.original,  # правильное слово
                            transcript=error.transcript,  # что сказал Яндекс
                            context=other.context,
                            marker_pos=other.marker_pos,
                            similarity=sim,
                            phonetic_similarity=sim,
                            is_yandex_error=False
                        ))

                        # 2. deletion: orig_word пропущено
                        fixed_errors.append(Error(
                            type='deletion',
                            time=error.time,
                            original=error.original,
                            context=error.context,
                            marker_pos=error.marker_pos
                        ))

                        skip_indices.add(orig_idx)
                        skip_indices.add(other_idx)
                        break

                elif other.type == 'insertion':
                    ins_word = other.transcript.lower()
                    # Проверяем: orig_word похоже на ins_word?
                    sim = fuzz.ratio(orig_word, ins_word) / 100.0

                    # Аналогичные проверки для insertion
                    len_diff = abs(len(orig_word) - len(ins_word))
                    min_len = min(len(orig_word), len(ins_word))

                    if (sim > 0.7 and
                        orig_word != trans_word and
                        ins_word != trans_word and
                        min_len > 2 and
                        len_diff <= 2):
                        # Аналогично: ins_word должно было сопоставиться с orig_word

                        fixed_errors.append(Error(
                            type='substitution',
                            time=other.time,
                            time_end=getattr(other, 'time_end', other.time),
                            original=error.original,
                            transcript=other.transcript,
                            context=error.context,
                            marker_pos=error.marker_pos,
                            similarity=sim,
                            phonetic_similarity=sim,
                            is_yandex_error=False
                        ))

                        fixed_errors.append(Error(
                            type='insertion',
                            time=error.time,
                            transcript=error.transcript,
                            context=other.context,
                            transcript_context=getattr(other, 'transcript_context', ''),
                            marker_pos=getattr(other, 'marker_pos', -1)
                        ))

                        skip_indices.add(orig_idx)
                        skip_indices.add(other_idx)
                        break

        # Если не исправили, добавляем как есть
        if orig_idx not in skip_indices:
            fixed_errors.append(error)

    return fixed_errors


# =============================================================================
# ПОСЕГМЕНТНОЕ СРАВНЕНИЕ (v9.2)
# =============================================================================

def compare_segment(
    orig_segment: List[str],
    trans_segment: List[str],
    orig_offset: int,
    trans_offset: int
) -> list:
    """
    Сравнивает один сегмент и возвращает opcodes с глобальными индексами.

    Args:
        orig_segment: нормализованные слова оригинала в сегменте
        trans_segment: нормализованные слова транскрипции в сегменте
        orig_offset: смещение начала сегмента в оригинале
        trans_offset: смещение начала сегмента в транскрипции

    Returns:
        Список opcodes с глобальными индексами
    """
    if not orig_segment and not trans_segment:
        return []

    matcher = SequenceMatcher(None, orig_segment, trans_segment)
    segment_opcodes = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        # Корректируем индексы на глобальные
        global_opcode = (
            tag,
            i1 + orig_offset,
            i2 + orig_offset,
            j1 + trans_offset,
            j2 + trans_offset
        )
        segment_opcodes.append(global_opcode)

    return segment_opcodes


def compare_with_segments(
    orig_norm: List[str],
    trans_norm: List[str],
    segments: list
) -> Tuple[list, float]:
    """
    Выполняет посегментное сравнение текстов.

    При посегментном сравнении SequenceMatcher работает с меньшими
    фрагментами текста, что даёт более точное выравнивание.

    Args:
        orig_norm: нормализованные слова оригинала
        trans_norm: нормализованные слова транскрипции
        segments: список Segment объектов от AlignmentManager

    Returns:
        (all_opcodes, average_ratio) — объединённые opcodes и средняя схожесть
    """
    all_opcodes = []
    total_ratio = 0.0
    segment_count = 0

    for segment in segments:
        # Извлекаем слова сегмента
        seg_orig = orig_norm[segment.orig_start:segment.orig_end]
        seg_trans = trans_norm[segment.trans_start:segment.trans_end]

        if not seg_orig and not seg_trans:
            continue

        # Сравниваем сегмент
        segment_opcodes = compare_segment(
            seg_orig, seg_trans,
            segment.orig_start, segment.trans_start
        )
        all_opcodes.extend(segment_opcodes)

        # Считаем схожесть сегмента
        if seg_orig or seg_trans:
            matcher = SequenceMatcher(None, seg_orig, seg_trans)
            total_ratio += matcher.ratio()
            segment_count += 1

    # Средняя схожесть по сегментам
    avg_ratio = total_ratio / segment_count if segment_count > 0 else 0.0

    return all_opcodes, avg_ratio


# =============================================================================
# ДЕТЕКТОР ТРАНСПОЗИЦИЙ (ПЕРЕСТАНОВОК СЛОВ)
# =============================================================================

def detect_transpositions_in_opcodes(opcodes: list, orig_norm: list, trans_norm: list,
                                      original: List[Word], transcript: List[Word]) -> Tuple[List[Error], set]:
    """
    Обнаруживает транспозиции (перестановки соседних слов) в opcodes.

    Паттерн транспозиции в SequenceMatcher:
    - INSERT X, EQUAL Y, DELETE X  означает Y X → X Y (слова поменялись местами)
    - DELETE X, EQUAL Y, INSERT X  означает X Y → Y X

    Args:
        opcodes: список операций от SequenceMatcher
        orig_norm: нормализованные слова оригинала
        trans_norm: нормализованные слова транскрипции
        original: полный список слов оригинала (для контекста)
        transcript: полный список слов транскрипции (для времени)

    Returns:
        (список ошибок transposition, множество индексов обработанных opcodes)
    """
    errors = []
    processed_indices = set()

    for i in range(len(opcodes) - 2):
        if i in processed_indices:
            continue

        op1 = opcodes[i]
        op2 = opcodes[i + 1]
        op3 = opcodes[i + 2]

        # Паттерн 1: INSERT, EQUAL, DELETE
        if op1[0] == 'insert' and op2[0] == 'equal' and op3[0] == 'delete':
            # op1: insert (i1, i1, j1, j2) - вставлено trans[j1:j2]
            # op2: equal (i1, i2, j2, j3) - совпадает orig[i1:i2] = trans[j2:j3]
            # op3: delete (i2, i3, j3, j3) - удалено orig[i2:i3]

            inserted = trans_norm[op1[3]:op1[4]]
            deleted = orig_norm[op3[1]:op3[2]]
            equal_word = orig_norm[op2[1]:op2[2]]

            # Проверяем: вставленное = удалённому? Это транспозиция!
            if inserted == deleted and len(inserted) == 1 and len(equal_word) == 1:
                word_a = equal_word[0]  # Первое слово в оригинале (которое осталось)
                word_b = inserted[0]    # Второе слово (которое переместилось)

                # Позиции для времени
                orig_pos_a = op2[1]  # Позиция word_a в оригинале
                orig_pos_b = op3[1]  # Позиция word_b в оригинале
                trans_pos_a = op1[3]  # Позиция word_b в транскрипции (где оно теперь)
                trans_pos_b = op2[3]  # Позиция word_a в транскрипции

                # Время из транскрипции
                time_a = transcript[trans_pos_a].time_start if trans_pos_a < len(transcript) else 0
                time_b = transcript[trans_pos_b].time_start if trans_pos_b < len(transcript) else 0

                # В оригинале: word_a word_b
                # В транскрипции: word_b word_a
                # Создаём одну ошибку типа transposition
                context, marker_pos = get_context(original, orig_pos_a)
                errors.append(Error(
                    type='transposition',
                    time=time_a,
                    time_end=time_b,
                    original=f"{word_a} {word_b}",  # Порядок в оригинале
                    transcript=f"{word_b} {word_a}",  # Порядок в транскрипции
                    context=context,
                    marker_pos=marker_pos,
                    is_yandex_error=False
                ))

                processed_indices.add(i)
                processed_indices.add(i + 1)
                processed_indices.add(i + 2)

        # Паттерн 2: DELETE, EQUAL, INSERT
        elif op1[0] == 'delete' and op2[0] == 'equal' and op3[0] == 'insert':
            deleted = orig_norm[op1[1]:op1[2]]
            inserted = trans_norm[op3[3]:op3[4]]
            equal_word = orig_norm[op2[1]:op2[2]]

            if deleted == inserted and len(deleted) == 1 and len(equal_word) == 1:
                word_a = deleted[0]     # Первое слово в оригинале
                word_b = equal_word[0]  # Второе слово в оригинале

                orig_pos_a = op1[1]
                trans_pos_b = op2[3]
                trans_pos_a = op3[3]

                time_b = transcript[trans_pos_b].time_start if trans_pos_b < len(transcript) else 0
                time_a = transcript[trans_pos_a].time_start if trans_pos_a < len(transcript) else 0

                # В оригинале: word_a word_b
                # В транскрипции: word_b word_a
                context, marker_pos = get_context(original, orig_pos_a)
                errors.append(Error(
                    type='transposition',
                    time=time_b,
                    time_end=time_a,
                    original=f"{word_a} {word_b}",
                    transcript=f"{word_b} {word_a}",
                    context=context,
                    marker_pos=marker_pos,
                    is_yandex_error=False
                ))

                processed_indices.add(i)
                processed_indices.add(i + 1)
                processed_indices.add(i + 2)

    return errors, processed_indices


# =============================================================================
# ГЛАВНАЯ ФУНКЦИЯ
# =============================================================================

def smart_compare(transcript_path: str, original_path: str,
                  audio_path: str = None, threshold: float = None,
                  output_path: str = None, phantom_seconds: float = None,
                  force: bool = False) -> dict:
    """
    Умное сравнение транскрипции с оригиналом.

    Использует SequenceMatcher для оптимального выравнивания последовательностей.

    Args:
        transcript_path: JSON файл с транскрипцией Яндекса
        original_path: текстовый файл с оригиналом
        audio_path: путь к аудио (для отчёта)
        threshold: порог схожести (0-1). По умолчанию из SmartCompareConfig
        output_path: путь для сохранения отчёта
        phantom_seconds: пропустить первые N секунд (None = автоопределение)
        force: перезаписать существующий файл без предупреждения

    Returns:
        dict с результатами
    """
    from difflib import SequenceMatcher

    # Значения по умолчанию из конфигурации
    if threshold is None:
        threshold = SmartCompareConfig.THRESHOLD if HAS_CONFIG else 0.7
    if phantom_seconds is None:
        phantom_seconds = SmartCompareConfig.PHANTOM_SECONDS if HAS_CONFIG else -1

    print(f"\n{'='*60}")
    print(f"  Умное сравнение транскрипции v{VERSION}")
    print(f"{'='*60}")

    # Автоопределение фантома (метаданных в начале)
    if phantom_seconds < 0:
        phantom_seconds = detect_phantom_seconds(transcript_path, original_path)
        if phantom_seconds > 0:
            print(f"\n  Автоопределение метаданных: {phantom_seconds:.1f} сек")
        else:
            phantom_seconds = 0

    # Парсим файлы
    print(f"\n  Загрузка транскрипции: {transcript_path}")
    transcript = parse_yandex_transcription(transcript_path, phantom_seconds=phantom_seconds)
    print(f"    Слов: {len(transcript)}")
    if phantom_seconds > 0:
        print(f"    (пропущено первые {phantom_seconds:.1f} сек — метаданные)")

    print(f"\n  Загрузка оригинала: {original_path}")
    original = parse_original_text(original_path)
    print(f"    Слов: {len(original)}")

    # v7.0: Склейка разбитых слов в транскрипции
    stitcher = ComparisonStitcher(original)
    transcript_before = len(transcript)
    transcript = stitcher.pre_stitch_audio_words(transcript)
    transcript = stitcher.normalize_hyphens(transcript)
    stitched_count = transcript_before - len(transcript)
    if stitched_count > 0:
        print(f"    Склеено слов: {stitched_count} (было {transcript_before}, стало {len(transcript)})")

    # Нормализуем списки слов для сравнения
    trans_norm = [w.normalized for w in transcript]
    orig_norm = [w.normalized for w in original]

    # v10.0: Поиск якорей для посегментного выравнивания
    anchors = []
    segments = []
    if HAS_ALIGNMENT_MANAGER:
        print(f"\n  Поиск якорей (AlignmentManager v1.1)...")
        alignment_mgr = AlignmentManager()
        anchors = alignment_mgr.find_anchor_points(orig_norm, trans_norm)
        segments = alignment_mgr.segment_by_anchors(len(orig_norm), len(trans_norm))
        print(f"    Найдено якорей: {len(anchors)}")
        print(f"    Сегментов: {len(segments)}")

        # v10.0: Разбиваем большие сегменты суб-якорями
        if segments:
            sizes = [s.orig_end - s.orig_start for s in segments]
            max_size = max(sizes)
            if max_size > 100:
                print(f"    Найден большой сегмент ({max_size} слов), ищем суб-якоря...")
                anchors, segments = alignment_mgr.refine_large_segments(orig_norm, trans_norm)
                print(f"    После уточнения: якорей={len(anchors)}, сегментов={len(segments)}")

            sizes = [s.orig_end - s.orig_start for s in segments]
            print(f"    Размер сегментов: avg={sum(sizes)/len(sizes):.1f}, max={max(sizes)}")

    # v10.0: Посегментное выравнивание через якоря
    print(f"\n  Выравнивание последовательностей...")
    if HAS_ALIGNMENT_MANAGER and anchors and segments:
        # Используем посегментное сравнение
        opcodes, ratio = compare_with_anchors(
            orig_norm, trans_norm, anchors, segments,
            original_words=original, transcript_words=transcript
        )
        print(f"    Режим: ПОСЕГМЕНТНЫЙ (якорей: {len(anchors)}, сегментов: {len(segments)})")
    else:
        # Fallback: полнотекстовое сравнение (если нет якорей)
        matcher = SequenceMatcher(None, orig_norm, trans_norm)
        ratio = matcher.ratio()
        opcodes = list(matcher.get_opcodes())
        print(f"    Режим: ПОЛНОТЕКСТОВЫЙ (нет якорей)")
    print(f"    Схожесть: {ratio*100:.1f}%")

    # Детектируем транспозиции
    transposition_errors, processed_indices = detect_transpositions_in_opcodes(
        opcodes, orig_norm, trans_norm, original, transcript
    )

    # Собираем ошибки
    all_errors = list(transposition_errors)

    for idx, (tag, i1, i2, j1, j2) in enumerate(opcodes):
        # Пропускаем opcodes, обработанные как транспозиции
        if idx in processed_indices:
            continue
        if tag == 'equal':
            continue

        elif tag == 'replace':
            # Замена: слова orig[i1:i2] заменены на trans[j1:j2]
            for k in range(max(i2 - i1, j2 - j1)):
                orig_idx = i1 + k
                trans_idx = j1 + k

                orig_word = original[orig_idx] if orig_idx < i2 else None
                trans_word = transcript[trans_idx] if trans_idx < j2 else None

                time = trans_word.time_start if trans_word else 0

                if orig_word and trans_word:
                    # Вычисляем схожесть
                    if HAS_THEFUZZ:
                        sim = fuzz.ratio(orig_word.text, trans_word.text) / 100.0
                    else:
                        sim = 0.5
                    phon_sim = phonetic_similarity(orig_word.text, trans_word.text) / 100.0
                    combined = max(sim, phon_sim)

                    # v6.0: Проверяем smart_rules для раннего определения ложных срабатываний
                    # Это позволяет пометить некоторые ошибки как "yandex" уже на этапе сравнения
                    is_yandex = False
                    if HAS_SMART_RULES:
                        smart_result = get_smart_rules().is_false_positive(trans_word.text, orig_word.text)
                        if smart_result and smart_result.is_match and smart_result.confidence >= 0.9:
                            is_yandex = True  # Высокая уверенность — помечаем как ошибку Яндекса

                    # v8.0: Проверяем ScoringEngine для защиты имён и hard negatives
                    penalty_info = ""
                    if HAS_SCORING_ENGINE and not is_yandex:
                        confidence = trans_word.confidence if hasattr(trans_word, 'confidence') else 0.8
                        should_filter, reason = should_filter_by_score(
                            orig_word.text, trans_word.text,
                            sim, phon_sim, confidence
                        )
                        # Если ScoringEngine говорит НЕ фильтровать — это вероятно реальная ошибка
                        # Сохраняем информацию для отладки
                        penalty_info = reason

                    # v10.1: Проверяем WindowVerifier для контекстной верификации
                    # Если контекст вокруг ошибки идентичен (без учёта пробелов) — это технический шум
                    window_verified = False
                    if HAS_WINDOW_VERIFIER and not is_yandex:
                        # Берём контекст: 3 слова до + слово + 3 слова после
                        ctx_start = max(0, orig_idx - 3)
                        ctx_end = min(len(original), orig_idx + 4)
                        trans_ctx_start = max(0, trans_idx - 3)
                        trans_ctx_end = min(len(transcript), trans_idx + 4)

                        orig_context_words = ' '.join(w.text for w in original[ctx_start:ctx_end])
                        trans_context_words = ' '.join(w.text for w in transcript[trans_ctx_start:trans_ctx_end])

                        verification = verify_segment(orig_context_words, trans_context_words)
                        if verification.status == VerificationStatus.TECHNICAL_OK:
                            # Контекст почти идентичен — это технический шум
                            is_yandex = True
                            window_verified = True

                    # ВСЕ substitution ошибки по умолчанию — это ошибки чтеца
                    # Яндекс распознал то, что услышал. Если результат отличается от оригинала,
                    # значит чтец прочитал не то (или не так).
                    #
                    # Ошибки Яндекса будут отфильтрованы позже в golden_filter
                    # по специфическим паттернам (типичные ошибки распознавания)

                    context, marker_pos = get_context(original, orig_idx)
                    all_errors.append(Error(
                        type='substitution',
                        time=time,
                        time_end=trans_word.time_end,
                        original=orig_word.text,
                        transcript=trans_word.text,
                        context=context,
                        marker_pos=marker_pos,
                        similarity=sim,
                        phonetic_similarity=phon_sim,
                        is_yandex_error=is_yandex
                    ))
                elif orig_word and not trans_word:
                    # Для deletion в replace блоке: берём время последнего слова блока транскрипции
                    del_time = time
                    if del_time == 0 and j2 > 0 and j2 <= len(transcript):
                        del_time = transcript[j2 - 1].time_end
                    elif del_time == 0 and j1 > 0:
                        del_time = transcript[j1 - 1].time_end
                    context, marker_pos = get_context(original, orig_idx)
                    all_errors.append(Error(
                        type='deletion',
                        time=del_time,
                        original=orig_word.text,
                        context=context,
                        marker_pos=marker_pos
                    ))
                elif trans_word and not orig_word:
                    # Для insertion: контекст из оригинала с маркером
                    orig_context_pos = min(i1 + (trans_idx - j1), len(original) - 1) if original else 0
                    if original:
                        context, marker_pos = get_context_with_marker(original, orig_context_pos)
                    else:
                        context, marker_pos = "", -1
                    all_errors.append(Error(
                        type='insertion',
                        time=time,
                        transcript=trans_word.text,
                        context=context,
                        transcript_context=get_context_from_transcript(transcript, trans_idx),
                        marker_pos=marker_pos
                    ))

        elif tag == 'delete':
            # Удаление: слова orig[i1:i2] отсутствуют
            # Вычисляем примерное время на основе соседних слов транскрипции
            # Берём время из ближайшего слова транскрипции
            time_estimate = 0.0
            if j1 < len(transcript):
                # Есть слово после удалённых — берём его время
                time_estimate = transcript[j1].time_start
            elif j1 > 0:
                # Есть слово до удалённых — берём его конечное время
                time_estimate = transcript[j1 - 1].time_end
            elif len(transcript) > 0:
                # Fallback: берём время последнего слова транскрипции
                time_estimate = transcript[-1].time_end

            for k in range(i1, i2):
                context, marker_pos = get_context(original, k)
                all_errors.append(Error(
                    type='deletion',
                    time=time_estimate,
                    original=original[k].text,
                    context=context,
                    marker_pos=marker_pos
                ))

        elif tag == 'insert':
            # Вставка: слова trans[j1:j2] лишние
            # Контекст из оригинала с маркером
            if i1 < len(original):
                context, marker_pos = get_context_with_marker(original, i1)
            else:
                context, marker_pos = "", -1
            for k in range(j1, j2):
                all_errors.append(Error(
                    type='insertion',
                    time=transcript[k].time_start,
                    transcript=transcript[k].text,
                    context=context,
                    transcript_context=get_context_from_transcript(transcript, k),
                    marker_pos=marker_pos
                ))

    # Пост-обработка: исправляем неправильные сопоставления SequenceMatcher
    # Например: "рагидон"→"и" + deletion "рагедон" → deletion "и" + substitution "рагидон"→"рагедон"
    all_errors = fix_misaligned_errors(all_errors)

    # Разделяем на типы
    yandex_errors = [e for e in all_errors if e.is_yandex_error]
    reader_errors = [e for e in all_errors if not e.is_yandex_error]

    print(f"\n  Результат:")
    print(f"    Всего несовпадений: {len(all_errors)}")
    print(f"    Ошибки Яндекса (фонетические): {len(yandex_errors)}")
    print(f"    Ошибки чтеца (реальные): {len(reader_errors)}")

    # Формируем отчёт
    report = {
        'audio': audio_path or '',
        'original': original_path,
        'transcript': transcript_path,
        'threshold': threshold,
        'stats': {
            'original_words': len(original),
            'transcript_words': len(transcript),
            'similarity': ratio,
            'total_differences': len(all_errors),
            'yandex_errors': len(yandex_errors),
            'reader_errors': len(reader_errors),
        },
        'errors': [asdict(e) for e in reader_errors],  # Только ошибки чтеца
        'yandex_errors': [asdict(e) for e in yandex_errors],  # Для отладки
    }

    # Определяем выходной путь
    if output_path:
        out_file = Path(output_path)
    else:
        # Используем FileNaming для правильного имени
        if HAS_CONFIG:
            chapter_id = FileNaming.get_chapter_id(Path(transcript_path))
            out_file = Path(transcript_path).parent / FileNaming.build_filename(chapter_id, 'compared')
        else:
            out_file = Path(transcript_path).with_stem(
                Path(transcript_path).stem + '_compared'
            )

    # Проверяем существование файла
    if out_file.exists() and not force:
        if HAS_CONFIG:
            if not check_file_exists(out_file, action='skip'):
                return report  # Пропускаем, файл уже существует
        else:
            print(f"  ⚠ Файл уже существует: {out_file.name}")

    with open(out_file, 'w', encoding='utf-8') as f:
        json.dump(report, f, ensure_ascii=False, indent=2)

    print(f"\n  Отчёт сохранён: {out_file}")
    print(f"{'='*60}\n")

    return report


def main():
    # Получаем значения по умолчанию из конфигурации
    default_threshold = SmartCompareConfig.THRESHOLD if HAS_CONFIG else 0.7
    default_phantom = SmartCompareConfig.PHANTOM_SECONDS if HAS_CONFIG else -1

    parser = argparse.ArgumentParser(
        description='Умное сравнение транскрипции с оригиналом',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=f"""
Алгоритм:
  1. Поиск якорей (100% совпадений) для синхронизации
  2. Анализ серых зон между якорями
  3. Классификация ошибок по схожести:
     - Высокая схожесть → фонетическая ошибка Яндекса (игнорируем)
     - Низкая схожесть → реальная ошибка чтеца (в отчёт)

Примеры:
  python smart_compare.py транскрипция.json оригинал.txt
  python smart_compare.py транскрипция.json оригинал.docx --audio глава.mp3
  python smart_compare.py транскрипция.json оригинал.txt --threshold 0.6
  python smart_compare.py транскрипция.json оригинал.txt --force

Текущие настройки из config.py:
  threshold: {default_threshold}
  phantom_seconds: {default_phantom} (-1 = авто)
        """
    )
    parser.add_argument('transcript', help='JSON файл с транскрипцией Яндекса')
    parser.add_argument('original', help='Текстовый файл с оригиналом (TXT или DOCX)')
    parser.add_argument('--audio', '-a', help='Путь к аудиофайлу')
    parser.add_argument('--output', '-o', help='Путь для сохранения отчёта')
    parser.add_argument('--threshold', '-t', type=float, default=None,
                        help=f'Порог схожести 0-1 (по умолчанию: {default_threshold})')
    parser.add_argument('--phantom', '-p', type=float, default=None,
                        help=f'Пропустить первые N секунд (-1 = авто, по умолчанию: {default_phantom})')
    parser.add_argument('--force', action='store_true',
                        help='Перезаписать существующие файлы')

    args = parser.parse_args()

    try:
        smart_compare(
            args.transcript,
            args.original,
            audio_path=args.audio,
            threshold=args.threshold,
            output_path=args.output,
            phantom_seconds=args.phantom,
            force=args.force
        )
        print("✓ Готово!")
    except Exception as e:
        print(f"✗ Ошибка: {e}")
        import traceback
        traceback.print_exc()


if __name__ == '__main__':
    main()
